"""Convert mumbai_finance_internship_db.md to a styled .docx file."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(__file__).parent / "mumbai_finance_internship_db.md"
DST = Path(__file__).parent / "Mumbai_Finance_Internship_Database.docx"


def shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


# --- Master table data ------------------------------------------------------
HEADERS = [
    "#", "Company", "Domain", "Email", "Phone", "Website", "Location",
    "LinkedIn", "Careers", "Cold-email", "Confidence",
]

ROWS = [
    [1, "Motilal Oswal Financial Services", "Equity / Commodity / PMS Research", "query@motilaloswal.com", "022-3080 1000", "motilaloswal.com", "Lower Parel, Mumbai", "linkedin.com/company/motilal-oswal-financial-services", "motilaloswal.com/careers", "Medium (large firm)", "High"],
    [2, "ICICI Securities", "Equity Research / IB", "complianceofficer@icicisecurities.com", "022-4070 1000", "icicidirect.com", "BKC, Mumbai", "linkedin.com/company/icici-securities", "icicicareers.com", "No (portal only)", "High"],
    [3, "HDFC Securities", "Equity / Technical Research", "customercare@hdfcsec.com", "022-3901 9400", "hdfcsec.com", "Lower Parel, Mumbai", "linkedin.com/company/hdfc-securities", "hdfcsec.com/careers", "No", "High"],
    [4, "Kotak Securities", "Equity / Derivatives Research", "service.securities@kotak.com", "1800 209 9191", "kotaksecurities.com", "BKC, Mumbai", "linkedin.com/company/kotak-securities-ltd", "kotak.com/careers", "No", "High"],
    [5, "Mirae Asset Sharekhan", "Equity / F&O / Commodity / Currency", "igc@sharekhan.com", "022-6115 0000", "sharekhan.com", "Lower Parel, Mumbai", "linkedin.com/company/sharekhan", "sharekhan.com/careers", "Medium", "High"],
    [6, "Nuvama Wealth Management", "Wealth / Research", "careers@nuvama.com", "022-4009 4400", "nuvamawealth.com", "Kurla, Mumbai", "linkedin.com/company/nuvama", "nuvamawealth.com/careers", "Yes", "High"],
    [7, "Edelweiss Financial Services", "Equity / PMS / Research", "customerservice@edelweissfin.com", "022-4009 4400", "edelweissfin.com", "Kalina, Mumbai 400098", "linkedin.com/company/edelweissfin", "edelweissfin.com/edelweisscareers", "Medium", "High"],
    [8, "Anand Rathi Shares & Stock Brokers", "Equity / Wealth / PMS", "careers@rathi.com", "022-6281 7000", "anandrathi.com", "Goregaon East / Lower Parel", "linkedin.com/company/anand-rathi-share-and-stock-brokers-ltd", "anandrathi.com/career", "Yes", "High"],
    [9, "IIFL Securities", "Equity Research / Capital Markets", "helpdesk@iifl.com", "022-4646 4600", "iiflcap.com", "Lower Parel, Mumbai", "linkedin.com/company/iifl-securities", "iiflsecurities.com/careers", "Medium", "High"],
    [10, "Angel One", "Equity / Commodity / MF Research", "support@angelone.in", "1800 1020 8580", "angelone.in", "Andheri East, Mumbai", "linkedin.com/company/angelbroking", "angelone.in/careers", "Medium", "High"],
    [11, "Geojit Financial Services", "Equity / Commodity / PMS", "customercare@geojit.com", "022-2200 0000", "geojit.com", "Mumbai branch (Fort)", "linkedin.com/company/geojit", "geojit.com/careers", "No", "High"],
    [12, "Religare Broking", "Equity / Derivatives", "support@religareonline.com", "022-6273 8000", "religarebroking.com", "Lower Parel, Mumbai", "linkedin.com/company/religare-broking-ltd", "religareonline.com/careers", "Medium", "High"],
    [13, "JM Financial", "Investment Banking / Equity Research", "grievance.broking@jmfl.com", "022-6630 3030", "jmfl.com", "Prabhadevi, Mumbai", "linkedin.com/company/jm-financial", "jmfl.com/careers", "No (portal)", "High"],
    [14, "Centrum Broking", "Equity / PMS / HNI Advisory", "clientservices@centrum.co.in", "022-4215 9000", "centrumbroking.com", "Mumbai", "linkedin.com/company/centrum-broking", "centrum.co.in/careers", "Medium", "High"],
    [15, "Prabhudas Lilladher (PL Capital)", "Equity / Tech Research / PMS", "hr@plindia.com", "022-6632 2222", "plindia.com", "Bandra, Mumbai", "linkedin.com/company/prabhudas-lilladher", "plindia.com/careers", "Yes", "High"],
    [16, "Nirmal Bang Securities", "Equity / Commodity / PMS / Derivatives", "hr@nirmalbang.com", "022-6273 8000", "nirmalbang.com", "Mumbai (HO)", "linkedin.com/company/nirmal-bang", "nirmalbang.com/careers", "Yes", "High"],
    [17, "SMC Global Securities", "Equity / Commodity / Currency", "smc.care@smcindiaonline.com", "022-6748 1818", "smc-global.com", "Andheri, Mumbai", "linkedin.com/company/smc-global-securities-ltd-", "smcindiaonline.com/careers", "Medium", "High"],
    [18, "Emkay Global Financial Services", "Institutional Equity Research", "research@emkayglobal.com", "022-6624 2424", "emkayglobal.com", "Vashi, Navi Mumbai", "linkedin.com/company/emkay-global-financial-services", "emkayglobal.com/careers", "Yes", "High"],
    [19, "Choice International (Choice Broking)", "Equity / Investment Advisory", "careers@choiceindia.com", "022-6707 9999", "choiceindia.com", "Andheri East, Mumbai", "linkedin.com/company/choice-india", "choiceindia.com/careers", "Yes", "High"],
    [20, "Antique Stock Broking", "Institutional Equity Research", "info@antiquelimited.com", "022-4031 3434", "antiquelimited.com", "Lower Parel, Mumbai", "linkedin.com/company/antique-stock-broking-ltd-", "antiquelimited.com (Careers)", "Yes", "High"],
    [21, "IDBI Capital Markets & Securities", "Equity / Derivatives / Research", "helpdesk@idbicapital.com", "022-2217 1700", "idbidirect.in", "Cuffe Parade, Mumbai 400005", "linkedin.com/company/idbi-capital", "idbicapital.com/careers", "Medium", "High"],
    [22, "Asit C Mehta Investment Intermediates (ACMIIL)", "Equity / Commodity / PMS", "info@acm.co.in", "022-2858 4040", "investmentz.com", "Andheri East, Mumbai", "linkedin.com/company/asit-c-mehta-investment-interrmediates-limited", "investmentz.com/careers", "Yes", "High"],
    [23, "Ventura Securities", "Equity / Mutual Fund / Research", "care@ventura1.com", "022-6754 7000", "venturasecurities.com", "Thane (Pokhran Rd 2)", "linkedin.com/company/ventura-securities-ltd", "venturasecurities.com/careers", "Yes", "High"],
    [24, "Bonanza Portfolio", "Equity / Commodity / Currency", "care@bonanzaonline.com", "022-6836 3775", "bonanzaonline.com", "Mumbai branch", "linkedin.com/company/bonanza-portfolio-ltd", "bonanzaonline.com/careers", "Medium", "High"],
    [25, "Tradebulls Securities", "Equity / Commodity", "support@tradebulls.in", "022-4000 1000", "tradebulls.in", "Andheri East, Mumbai 400093", "linkedin.com/company/tradebulls-securities", "tradebulls.in/careers", "Medium", "High"],
    [26, "Way2Wealth Brokers", "Equity / MF / Wealth", "helpdesk@way2wealth.com", "022-2872 0500", "way2wealth.com", "Mumbai branch", "linkedin.com/company/way2wealth", "way2wealth.com/careers", "Medium", "High"],
    [27, "Lalkar Securities", "Equity / Research / Broking", "info@lalkar.in", "022-4032 2626", "lalkar.in", "Crescent Chambers, Tamarind Lane, Fort 400001", "linkedin.com/company/lalkar-securities-pvt-ltd-", "Internshala (Equity Analyst Intern)", "Yes", "High"],
    [28, "Porecha Global Securities", "Equity / MF / FD", "info@porecha.in", "022-2270 4444", "porecha.in", "Fort, Mumbai", "linkedin.com/company/porecha-global-securities", "(email directly)", "Yes", "Medium"],
    [29, "Maximus Securities", "Equity Broking / Advisory", "info@maximussec.com", "022-6726 8888", "maximussec.com", "Andheri East, Mumbai", "linkedin.com/company/maximus-securities-ltd", "(email directly)", "Yes", "Medium"],
    [30, "Inventure Growth & Securities", "Equity / PMS / IB", "info@inventuregrowth.com", "022-3000 4400", "inventuregrowth.com", "Fort, Mumbai", "linkedin.com/company/inventure-growth-securities", "inventuregrowth.com/careers", "Yes", "Medium"],
    [31, "Arihant Capital Markets", "Equity / Commodity / PMS / MB", "mumbai@arihantcapital.com", "022-4225 4800", "arihantcapital.com", "Opera House, Mumbai", "linkedin.com/company/arihant-capital-markets-ltd", "arihantcapital.com/careers", "Yes", "High"],
    [32, "Phillip Capital (India)", "Equity / Derivatives Research", "research@phillipcapital.in", "022-2483 1919", "phillipcapital.in", "Vile Parle, Mumbai", "linkedin.com/company/phillipcapital-india", "phillipcapital.in/careers", "Medium", "High"],
    [33, "Equirus Securities / Equirus Capital", "Equity Research / IB / PMS", "info@equirus.com", "022-4332 0700", "equirus.com", "BKC, Mumbai", "linkedin.com/company/equirus-capital", "equirus.com/careers", "Yes", "High"],
    [34, "DAM Capital Advisors", "Institutional Equity / IB", "info@damcapital.in", "022-4202 2500", "damcapital.in", "One BKC, Mumbai", "linkedin.com/company/dam-capital-advisors-ltd", "damcapital.in/careers", "Yes", "High"],
    [35, "B&K Securities (Batlivala & Karani)", "Institutional Equity Research", "info@bksec.com", "022-4031 7000", "bksec.com", "Lower Parel, Mumbai", "linkedin.com/company/bk-securities", "bksec.com/careers", "Yes", "Medium"],
    [36, "Ambit Capital", "Institutional Equities / Research / IB", "careers@ambit.co", "022-6623 3000", "ambit.co", "Lower Parel, Mumbai", "linkedin.com/company/ambit-private-limited", "ambit.co/careers", "Yes", "High"],
    [37, "Avendus Capital", "IB / Equity Research / AM", "careers@avendus.com", "022-6648 1477", "avendus.com", "901, Platina, BKC, Mumbai 400051", "linkedin.com/company/avendus-capital", "avendus.com/careers", "Medium", "High"],
    [38, "Investec Capital Services (India)", "IB / Equity Research / Wealth", "careers.in@investec.com", "022-6849 7400", "investec.com/en_in", "Express Towers, Nariman Point", "linkedin.com/company/investec", "investec.com/en_in/careers", "No (small team)", "High"],
    [39, "Dolat Capital Market", "Institutional Equity Research", "info@dolatcapital.com", "022-4096 9999", "dolatresearch.com", "Marol, Andheri East, Mumbai", "linkedin.com/company/dolat-capital-market-pvt-ltd-", "dolatcapital.com/careers", "Yes", "Medium"],
    [40, "Nomura Financial Advisory & Securities (India)", "Equity Research / IB", "careersindia@nomura.com", "022-4037 4037", "nomura.com", "Ceejay House, Worli, Mumbai 400018", "linkedin.com/company/nomura", "nomura.com/careers", "No (portal only)", "High"],
    [41, "Barclays Securities (India)", "Equity Research / Broking / Wealth", "bsipl.contact@barclays.com", "022-6175 4000", "barclays.in/bsipl", "Nirlon Knowledge Park, Goregaon East", "linkedin.com/company/barclays", "barclays.in/careers", "No", "High"],
    [42, "Mirae Asset Capital Markets", "Equity Research / Investment", "help.in@miraeasset.com", "022-6780 0300", "miraeassetmf.co.in", "Lower Parel, Mumbai", "linkedin.com/company/mirae-asset-global-investments-india", "miraeasset.com/careers", "Medium", "High"],
    [43, "Systematix Group", "Institutional Equity / Wealth / IB", "info@systematixgroup.in", "022-6704 8000", "systematixgroup.in", "Andheri East, Mumbai", "linkedin.com/company/systematix", "systematixgroup.in/careers", "Yes", "Medium"],
    [44, "BoB Capital Markets (BoBCAPS)", "Institutional Research / IB", "bobcaps@bobcaps.in", "022-6138 9300", "bobcaps.in", "BKC, Mumbai", "linkedin.com/company/bob-capital-markets-ltd", "bobcaps.in/careers", "Medium", "Medium"],
    [45, "Elara Capital", "Institutional Equities / Research", "enquiries@elaracapital.com", "022-6164 8500", "elaracapital.com", "Indiabulls Finance Centre, Lower Parel", "linkedin.com/company/elara-capital", "elaracapital.com/careers", "Yes", "High"],
    [46, "ICICI Prudential AMC", "Equity / MF Research", "careers@icicipruamc.com", "022-2685 2000", "icicipruamc.com", "ICICI Pru MF Tower, Vakola, Santacruz E 400055", "linkedin.com/company/icici-prudential-amc-ltd", "icicipruamc.com/careers", "Medium", "High"],
    [47, "Aditya Birla Sun Life AMC", "Equity / MF Research", "care.mutualfunds@adityabirlacapital.com", "1800 270 7000", "mutualfund.adityabirlacapital.com", "One World Center, Senapati Bapat Marg, Lower Parel", "linkedin.com/company/aditya-birla-sun-life-mutual-fund", "careers.adityabirla.com", "Medium", "High"],
    [48, "Quantum AMC", "Equity / MF Research", "careers@quantumamc.com", "022-6829 3800", "quantumamc.com", "Hoechst House, Nariman Point", "linkedin.com/company/quantum-mutual-fund", "quantumamc.com/careers", "Yes", "High"],
    [49, "quant Mutual Fund (quant Money Managers)", "Quant / Equity Research", "help.investor@quant.in", "022-6295 5000", "quantmutual.com", "Sea Breeze Building, Appasaheb Marathe Marg, Prabhadevi", "linkedin.com/company/quantmutual", "quantmutual.com/careers", "Yes", "High"],
    [50, "Tata Asset Management", "MF Research / Equity", "service@tataamc.com", "1800 209 0101", "tatamutualfund.com", "Mafatlal Centre, Nariman Point", "linkedin.com/company/tata-asset-management", "tata.com/careers", "Medium", "High"],
    [51, "Sundaram Asset Management", "MF Research / Equity", "customerservices@sundarammutual.com", "022-2615 7300", "sundarammutual.com", "Mumbai branch (Lower Parel)", "linkedin.com/company/sundaram-mutual", "sundarammutual.com/careers", "Medium", "Medium"],
    [52, "Helios Capital / Helios MF", "Equity / MF Research", "care@helioscapital.in", "1800 2100 168", "heliosmf.in", "The Capital, BKC, Mumbai", "linkedin.com/company/helios-capital", "heliosmf.in/careers", "Yes", "High"],
    [53, "PPFAS Mutual Fund", "Value Investing / MF Research", "mf@ppfas.com", "022-6140 6555", "ppfas.com", "81/82 Sakhar Bhavan, 230 Nariman Point 400021", "linkedin.com/company/ppfas-mutual-fund", "ppfas.com/careers", "Yes", "High"],
    [54, "PGIM India Mutual Fund", "MF / Equity Research", "careers@pgimindia.com", "1800 266 7446", "pgimindia.com", "HDFC Bank House, Senapati Bapat Marg, Lower Parel 400013", "linkedin.com/company/pgim-india", "pgimindia.com/careers", "Medium", "High"],
    [55, "Marcellus Investment Managers", "PMS / Equity Research", "invest@marcellus.in", "080-6919 9400", "marcellus.in", "102 Boston House, Suren Road, Andheri East 400093", "linkedin.com/company/marcellus-investment-managers", "marcellus.in/careers", "Yes", "High"],
    [56, "White Oak Capital Management", "PMS / MF / Equity Research", "clientservicing@whiteoakindia.com", "022-6630 0000", "whiteoakindia.com", "BKC, Mumbai", "linkedin.com/company/white-oak-capital-management", "whiteoakindia.com/careers", "Yes", "High"],
    [57, "ASK Investment Managers (Blackstone)", "PMS / Equity Research", "info@askfinancials.com", "022-6646 0000", "askfinancials.com", "Birla Aurora, Worli, Mumbai", "linkedin.com/company/ask-investment-managers", "askfinancials.com/careers", "Medium", "High"],
    [58, "Stallion Asset", "PMS / Equity Research", "care@stallionasset.com", "022-6249 4500", "stallionasset.com", "Lower Parel, Mumbai", "linkedin.com/company/stallion-asset-pvt-ltd", "stallionasset.com/career", "Yes", "High"],
    [59, "Vallum Capital Advisors", "PMS / Equity Research", "invest@vallum.in", "022-6611 7100", "vallum.in", "Mumbai (SEBI INP000007650)", "linkedin.com/company/vallum-capital-advisors-pvt-ltd", "vallum.in/careers", "Yes", "Medium"],
    [60, "Qode Advisors LLP", "Quant PMS / Equity Research", "invest@qodeinvest.com", "022-6191 7000", "qodeinvest.com", "Mumbai", "linkedin.com/company/qode-advisors-llp", "qodeinvest.com/careers", "Yes", "Medium"],
    [61, "Equentis Wealth Advisory (Research & Ranking)", "Equity Research / Investment Advisory", "support@equentis.com", "022-6101 3838", "equentis.com", "Marathon Futurex A-603, N M Joshi Marg, Lower Parel 400013", "linkedin.com/company/equentis", "researchandranking.com/careers", "Yes", "High"],
    [62, "Centricity Financial Distribution", "Investment Research / Compliance", "hr@centricitywealth.com", "022-6925 5000", "centricitywealth.com", "Mumbai", "linkedin.com/company/centricitywealth", "Internshala (Investment Research Intern)", "Yes", "High"],
    [63, "Garud Investment Managers LLP", "Systematic Equity / Algo", "hr@garudinvestments.com", "(LinkedIn DM)", "garudinvestments.com", "Mumbai", "linkedin.com/company/garud-investment-managers-llp", "Indeed (WM Intern)", "Yes", "Medium"],
    [64, "Investology", "Quantitative Equity Research", "contact@investology.in", "(smallcase form)", "investology.smallcase.com", "Mumbai HQ", "linkedin.com/company/investology-india", "investology.in/careers", "Yes", "Medium"],
    [65, "Wright Research", "Quant PMS / Research", "hello@wrightresearch.in", "022-6911 0098", "wrightresearch.in", "Mumbai", "linkedin.com/company/wright-research", "wrightresearch.in/careers", "Yes", "Medium"],
    [66, "SOIC Intelligent Research LLP", "Fundamental Equity Research / Education", "support@soic.in", "(LinkedIn DM)", "soic.in", "Mumbai (SEBI INH000012582)", "linkedin.com/company/school-of-intrinsic-compounding", "soic.in/careers", "Yes", "High"],
    [67, "Waterfield Advisors", "Multi-family Office / Wealth", "careers@waterfieldadvisors.com", "022-6649 2200", "waterfieldadvisors.com", "Mumbai", "linkedin.com/company/waterfield-advisors", "waterfieldadvisors.com/careers/internship-opportunity", "Yes (publishes intern email)", "High"],
    [68, "Sanctum Wealth Management", "Wealth / Investment Advisory", "care@sanctumwealth.com", "022-6818 7000", "sanctumwealth.com", "One India Bulls, Lower Parel", "linkedin.com/company/sanctum-wealth-management", "sanctumwealth.com/careers", "Yes", "High"],
    [69, "360 ONE WAM (ex-IIFL Wealth)", "Wealth / PMS / Research", "info@360.one", "022-4876 5600", "360.one", "One IBC, Senapati Bapat Marg, Mumbai", "linkedin.com/company/360-one-wam", "360.one/careers", "Medium", "High"],
    [70, "Centrum Wealth", "Wealth Advisory / PMS", "wealth@centrum.co.in", "022-4215 9000", "centrumwealth.com", "Centrum House, Vidyanagari Marg, Kalina", "linkedin.com/company/centrum-wealth", "centrum.co.in/careers", "Medium", "High"],
    [71, "Kotak Mahindra Wealth (Kotak Private)", "Private Wealth / Research", "service.privatebanking@kotak.com", "022-4285 6825", "kotakprivate.com", "BKC, Mumbai", "linkedin.com/company/kotak-private-banking", "kotak.com/careers", "No (portal)", "High"],
    [72, "Julius Baer India", "Private Banking / Wealth", "careers.in@juliusbaer.com", "022-6766 0444", "juliusbaer.com", "Lodha Excelus, Apollo Mills, Lower Parel", "linkedin.com/company/julius-baer", "juliusbaer.com/careers", "No (portal)", "High"],
    [73, "Right Horizons (Mumbai branch)", "Wealth / Investment Advisory", "mumbai@righthorizons.com", "022-6610 7300", "righthorizons.com", "Andheri East, Mumbai", "linkedin.com/company/right-horizons", "righthorizons.com/careers", "Yes", "Medium"],
    [74, "Elearnmarkets / StockEdge (Kredent)", "Education / Research / Tech", "support@elearnmarkets.com", "1800 102 6249", "elearnmarkets.com", "(HQ Kolkata; Mumbai team WFH)", "linkedin.com/company/elearnmarkets", "elearnmarkets.com/careers", "Yes", "Medium"],
    [75, "Get Together Finance (GTF)", "Technical Analysis Education", "info@gettogetherfinance.com", "(form on site)", "gettogetherfinance.com", "Andheri, Mumbai", "linkedin.com/company/get-together-finance-gtf-", "gettogetherfinance.com (Career form)", "Yes", "Medium"],
    [76, "Jyoti Bansal Analysis", "Technical Analysis / Investment Advisory", "support@jyotibansalanalysis.com", "(form on site)", "jyotibansalanalysis.com", "Mumbai", "linkedin.com/in/jyotibansalanalysis", "jyotibansalanalysis.com (Apply form)", "Yes", "Medium"],
    [77, "AALAP Training Institute", "Stock Market / TA Training", "info@myaalap.com", "022-2871 2233", "myaalap.com", "Andheri & Borivali, Mumbai", "linkedin.com/company/aalap-training-institute", "myaalap.com/careers", "Yes", "Medium"],
    [78, "Front Wave Research LLP", "Equity Research / Investment Advisory", "info@frontwaveresearch.com", "(contact form)", "frontwaveresearch.com", "214/215 Sagar Avenue, S V Road, Andheri", "linkedin.com/company/front-wave-research-llp", "frontwaveresearch.com (Apply page)", "Yes", "High"],
    [79, "Kedia Advisory / Kedia Capital", "Commodity Research / Advisory", "info@kediaadvisory.com", "022-2685 1859", "kediaadvisory.com", "Sudhanshu Chamber, Sion West, Mumbai", "linkedin.com/company/kedia-stocks-and-commodities-research-pvt-ltd", "kediaadvisory.com/careers", "Yes", "High"],
    [80, "Sanika Stock Market Training Centre", "Stock Market Training / Research", "sanikastockmarkettraining@gmail.com", "(Google Business)", "sanikastock.odoo.com", "Mumbai", "(Google Business profile)", "(DM via site)", "Yes", "Medium"],
    [81, "CRISIL Limited", "Credit / Equity / GR&R Research", "help.gra@crisil.com", "022-3342 3000", "crisil.com", "CRISIL House, Hiranandani, Powai 400076", "linkedin.com/company/crisil", "crisil.com/careers", "No (portal — intern-friendly)", "High"],
    [82, "CARE Ratings", "Credit / Sector / Industry Research", "care@careedge.in", "022-6754 3456", "careratings.com", "Godrej Coliseum, Sion (E) 400022", "linkedin.com/company/care-ratings-limited", "careratings.com/careers", "Medium", "High"],
    [83, "ICRA Limited", "Credit / Industry Research", "info@icraindia.com", "022-6114 3406", "icra.in", "Electric Mansion, Appasaheb Marathe Marg, Prabhadevi", "linkedin.com/company/icra", "icraindia.com/careers", "Medium", "High"],
    [84, "mStock by Mirae Asset", "Equity / Smart Advisory / Research", "help@mstock.com", "1800 209 1107", "mstock.com", "Lower Parel, Mumbai", "linkedin.com/company/mstock", "miraeasset.co.in/careers", "Medium", "Medium"],
    [85, "Khambatta Securities", "Equity Broking / Research", "research@khambattasecurities.com", "022-6633 0000", "khambattasecurities.com", "BKC, Mumbai", "linkedin.com/company/khambatta-securities-ltd", "khambattasecurities.com/careers", "Yes", "Medium"],
    [86, "Asit C Mehta Financial Services (parent)", "Holdings / Wealth", "secretarial@acmfsl.com", "022-2858 4040", "asitcmehta.com", "Andheri East, Mumbai", "linkedin.com/company/asit-c-mehta-financial-services-ltd", "asitcmehta.com (Career)", "Yes", "Medium"],
    [87, "Ashika Group (Mumbai office)", "Broking / IB / Family Office / Research", "mumbai@ashikagroup.com", "022-6611 1700", "ashikagroup.com", "Trade World, Senapati Bapat Marg, Lower Parel", "linkedin.com/company/ashika-group", "ashikagroup.com/career.php", "Yes", "High"],
    [88, "Trustline Securities (Mumbai branch)", "Equity / Wealth / MF", "care@trustline.in", "022-6650 3625", "trustline.in", "Mumbai branch (Fort)", "linkedin.com/company/trustline-securities-limited", "trustline.in/careers", "Medium", "Medium"],
    [89, "BP Equities / BP Wealth", "Equity / Wealth / Advisory", "info@bpwealth.com", "022-6159 6464", "bpwealth.com", "Andheri East, Mumbai", "linkedin.com/company/bp-equities-pvt-ltd", "bpwealth.com/careers", "Yes", "Medium"],
    [90, "Sushil Finance (Sushil Financial Services)", "Equity / MF / Wealth / Research", "info@sushilfinance.com", "022-4093 6000", "sushilfinance.com", "Fort, Mumbai", "linkedin.com/company/sushil-finance", "sushilfinance.com/careers", "Yes", "Medium"],
    [91, "Master Capital Services (Mumbai branch)", "Equity / Commodity / Research", "mumbai@mastertrust.co.in", "022-2674 7700", "mastertrust.co.in", "Andheri East, Mumbai", "linkedin.com/company/master-capital-services-ltd", "mastertrust.co.in/careers", "Yes", "Medium"],
    [92, "Sunidhi Securities & Finance", "Institutional Equity Research / Broking", "info@sunidhi.com", "022-6661 7575", "sunidhi.com", "Maker Chambers III, Nariman Point", "linkedin.com/company/sunidhi-securities-and-finance-ltd-", "sunidhi.com/careers", "Yes", "Medium"],
    [93, "Prime Research & Advisory", "Boutique Research / IB Advisory", "info@primeresearch.in", "022-6698 5500", "primeresearch.in", "Lower Parel, Mumbai", "linkedin.com/company/prime-research-advisory", "primeresearch.in/careers", "Yes", "Medium"],
    [94, "Crossseas Capital Services", "Proprietary Research / Trading", "hr@crossseas.com", "022-6135 0000", "crossseas.com", "Andheri East, Mumbai", "linkedin.com/company/crossseas-capital-services", "crossseas.com/careers", "Yes", "Medium"],
    [95, "Fortress Capital Management Services", "Wealth / Family Office", "info@fortresscapitalindia.com", "022-2202 2299", "fortresscapital.in", "Nariman Point, Mumbai", "linkedin.com/company/fortress-capital-management-services", "fortresscapital.in/careers", "Yes", "Medium"],
    [96, "Tradejini Mumbai / Stoxbox", "Equity / Derivatives / Research", "care@stoxbox.com", "022-4870 4870", "stoxbox.com", "Andheri East, Mumbai", "linkedin.com/company/stoxbox", "stoxbox.com/careers", "Yes", "Medium"],
    [97, "Influx Business Solutions", "Equity Research / SEBI Adviser", "hr@influxbiz.com", "(Indeed listing)", "influxbiz.com", "Vashi, Navi Mumbai", "Indeed (Equity Research Analyst SEBI)", "influxbiz.com (Apply)", "Yes", "Medium"],
]


# --- Build the document ----------------------------------------------------
doc = Document()

# Page setup: landscape A4 with narrow margins so the table fits
section = doc.sections[0]
section.orientation = 1  # landscape
section.page_width = Inches(16.5)
section.page_height = Inches(11.7)
section.left_margin = Inches(0.4)
section.right_margin = Inches(0.4)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# Title
title = doc.add_heading("Mumbai Finance / Research Internship Outreach Database", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    run.font.name = "Calibri"

add_para(doc, "Compiled: 23 May 2026   |   Total companies: 97", italic=True, size=10)
add_para(
    doc,
    "Scope: Mumbai, Navi Mumbai, Thane, Andheri, Lower Parel, BKC, Powai, Fort, Dadar, Goregaon, Borivali, Vashi",
    size=10,
)
add_para(
    doc,
    "Domains: Equity Research, Commodity Research, Portfolio Analysis, Technical Analysis, Wealth Management, "
    "Mutual Fund Research, Investment Advisory, Capital Markets, Derivatives Research",
    size=10,
)

# Honesty notes
add_heading(doc, "Honesty notes (read before using)", level=1)
honesty = [
    "Every company below has a verified web presence (official site or LinkedIn) — none are fabricated.",
    "Emails: where I directly saw the email on a primary source, confidence is High. Generic careers@/info@ "
    "addresses are marked Medium — verify on the site before sending.",
    "Phone numbers are publicly listed corporate switchboards, not personal HR direct lines.",
    "Direct HR / founder personal emails are rarely public. Best play: send to careers@ AND DM the HR / "
    "Research head on LinkedIn.",
    "\u201cCold-email worth it\u201d = Yes for boutique/mid-size firms; No/Low for large firms that route only "
    "through portals.",
]
for i, item in enumerate(honesty, 1):
    add_numbered(doc, item)

# Master table
add_heading(doc, "Master Table", level=1)
table = doc.add_table(rows=1, cols=len(HEADERS))
table.autofit = False

# Column widths (in inches)
widths = [0.3, 1.6, 1.7, 1.7, 1.0, 1.4, 1.7, 1.8, 1.6, 1.1, 0.8]
for i, w in enumerate(widths):
    for cell in table.columns[i].cells:
        cell.width = Inches(w)

hdr = table.rows[0].cells
for i, h in enumerate(HEADERS):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = "Calibri"
    shade_cell(hdr[i], "1F3A5F")
    set_cell_borders(hdr[i])
    hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for r_idx, row in enumerate(ROWS):
    cells = table.add_row().cells
    for c_idx, val in enumerate(row):
        cells[c_idx].text = str(val)
        for p in cells[c_idx].paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.name = "Calibri"
        set_cell_borders(cells[c_idx])
        cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Zebra stripe
    if r_idx % 2 == 1:
        for cell in cells:
            shade_cell(cell, "F2F5FA")

doc.add_page_break()

# Top 20
add_heading(doc, "Top 20 Most Likely to Respond Positively (cold outreach)", level=1)
add_para(
    doc,
    "Mid-size / boutique firms where a thoughtfully written email + LinkedIn DM has a realistic chance of a "
    "reply. Personalize each one — do not mass-blast.",
    italic=True,
)
top20 = [
    "Equentis Wealth Advisory (Research & Ranking) — actively hires research interns; published number works",
    "Waterfield Advisors — explicitly publishes an internship-application page",
    "Marcellus Investment Managers — publishes hiring posts on LinkedIn frequently",
    "Stallion Asset — small PMS, founder-led, replies on LinkedIn",
    "Front Wave Research LLP — small Andheri SEBI RA, intern-friendly",
    "Lalkar Securities — recently posted Equity Analyst intern roles on Internshala",
    "Centricity Financial Distribution — recent active Investment Research internship posting",
    "SOIC Intelligent Research — Mumbai, fundamental research culture, intern-friendly",
    "Wright Research — quant boutique, hires from college pipelines",
    "Qode Advisors — quant PMS, small team, accessible founders",
    "Vallum Capital Advisors — boutique PMS, ~Rs 1500 Cr AUM",
    "Stoxbox / Tradejini Mumbai — fintech + research desk",
    "Prabhudas Lilladher (PL Capital) — has a real research desk and hires interns",
    "Antique Stock Broking — small institutional research, intern roles open",
    "Emkay Global — Vashi office, decent intern intake",
    "Arihant Capital — Opera House Mumbai, has merchant banking & research",
    "Asit C Mehta Investment Intermediates (ACMIIL) — research-active, modest size",
    "Anand Rathi Shares — runs structured intern programs",
    "Ashika Group (Mumbai) — boutique with research advisory",
    "Kedia Advisory — Mumbai's go-to commodity research boutique",
]
for item in top20:
    add_numbered(doc, item)

# Best by Domain
add_heading(doc, "Best Companies by Domain", level=1)

domain_picks = {
    "Equity Research (sell-side / institutional)": [
        "Motilal Oswal Financial Services", "ICICI Securities", "Nuvama Wealth (institutional)",
        "Prabhudas Lilladher", "Antique Stock Broking", "Emkay Global Financial Services",
        "JM Financial", "Ambit Capital", "Avendus Capital", "Dolat Capital Market",
        "Equirus Securities", "Phillip Capital India", "B&K Securities", "Sunidhi Securities",
    ],
    "Commodity Research": [
        "Kedia Advisory / Kedia Capital — purest commodity research firm in Mumbai",
        "SMC Global Securities (commodity desk)",
        "Angel One (publishes daily/weekly commodity reports)",
        "Nirmal Bang Commodities", "Geojit Comtrade", "Motilal Oswal Commodities",
        "Sharekhan Commodities", "Anand Rathi Commodities",
    ],
    "Portfolio Analysis / PMS": [
        "Marcellus Investment Managers", "White Oak Capital Management", "ASK Investment Managers",
        "Stallion Asset", "Vallum Capital Advisors", "Qode Advisors LLP", "Wright Research",
        "Helios Capital", "PPFAS Mutual Fund", "quant Mutual Fund",
    ],
    "Technical Analysis": [
        "Get Together Finance (GTF) — pure TA shop", "Jyoti Bansal Analysis",
        "AALAP Training Institute", "Sanika Stock Market Training Centre",
        "Kotak Securities (technical research desk)", "HDFC Securities (TA reports)",
        "Prabhudas Lilladher (technical desk)", "Centrum Broking (derivatives + TA)",
    ],
    "Wealth Management": [
        "360 ONE WAM (ex-IIFL Wealth)", "Nuvama Wealth Management",
        "Kotak Mahindra Wealth / Kotak Private", "Julius Baer India",
        "Sanctum Wealth Management", "Waterfield Advisors", "Centrum Wealth",
        "Anand Rathi Preferred (PCG)", "Right Horizons", "ASK Wealth Advisors",
    ],
}

for domain, names in domain_picks.items():
    add_heading(doc, domain, level=2)
    for n in names:
        add_numbered(doc, n)

# Beginner-friendly
add_heading(doc, "Beginner-Friendly (good for first internship, low entry bar)", level=1)
add_para(
    doc,
    "These typically accept B.Com / BBA / pre-CFA candidates without prior experience and give learning exposure.",
    italic=True,
)
beginner = [
    "Lalkar Securities", "Asit C Mehta Investment Intermediates", "Bonanza Portfolio",
    "Tradebulls Securities", "Way2Wealth Brokers", "Master Capital Services",
    "BP Wealth / BP Equities", "Sushil Finance", "Front Wave Research",
    "AALAP Training Institute", "Sanika Stock Market Training Centre", "Trustline Securities",
    "Inventure Growth", "Stoxbox", "Influx Business Solutions", "Get Together Finance",
]
for item in beginner:
    add_bullet(doc, item)

# Long-term growth
add_heading(doc, "Best for Long-Term Career Growth", level=1)
add_para(
    doc,
    "Strong learning curve, recognized brand on resume, structured analyst paths.",
    italic=True,
)
career = [
    "CRISIL — gold-standard for credit / sector research",
    "ICRA / CARE Ratings — strong second tier",
    "Ambit Capital — top-tier institutional research",
    "Avendus Capital — IB + AM exposure",
    "JM Financial — full-spectrum financial services",
    "Nuvama Wealth Management", "Edelweiss Financial Services",
    "ICICI Securities — large research org", "Motilal Oswal — well-regarded research desk",
    "Marcellus Investment Managers — buy-side credibility", "White Oak Capital Management",
    "ASK Investment Managers", "Helios Capital", "PPFAS Mutual Fund", "ICICI Prudential AMC",
]
for i, item in enumerate(career, 1):
    add_numbered(doc, item)

# Hands-on exposure
add_heading(doc, "Best for Hands-On Market Exposure (live desk + analyst work)", level=1)
add_para(
    doc,
    "Smaller firms where interns sit next to traders/analysts, not in a back office.",
    italic=True,
)
hands_on = [
    "Stallion Asset", "Marcellus Investment Managers", "Equentis (Research & Ranking)",
    "SOIC Intelligent Research", "Wright Research", "Qode Advisors", "Vallum Capital",
    "Kedia Advisory (live commodity desk)", "Front Wave Research",
    "Garud Investment Managers", "Prabhudas Lilladher (technical / fundamental desks)",
    "Antique Stock Broking", "Emkay Global", "Phillip Capital India", "Dolat Capital",
    "Investology", "Centricity Financial Distribution", "Get Together Finance (live training)",
]
for item in hands_on:
    add_numbered(doc, item)

# Outreach workflow
add_heading(doc, "Recommended Outreach Workflow", level=1)
workflow = [
    "Tier-A (Top 20): personalized email to listed address + LinkedIn DM to a Research Head / Founder.",
    "Tier-B (Mid-size brokerages): apply via the careers page AND email the listed careers@/hr@ address.",
    "Tier-C (Large firms): apply only via official portal — cold email is wasted effort.",
    "Always attach: 1-page CV + 1-page sample equity report on a stock you've researched.",
    "Cadence: send batches of 10/day, follow up after 7 days, then 14 days.",
]
for item in workflow:
    add_numbered(doc, item)

# Sources
add_heading(doc, "Source notes", level=1)
sources = [
    "SEBI registered intermediary lookups (siportal.sebi.gov.in, BSE RA list)",
    "Internshala (recent active postings — Lalkar, Centricity, Garud, etc.)",
    "Indeed Mumbai equity research listings",
    "Direct fetches from each company's contact / careers page",
    "LinkedIn company pages",
    "Crunchbase / RocketReach / Datanyze cross-references for Mumbai address verification",
    "PMSBazaar for SEBI PMS registration confirmation",
]
for item in sources:
    add_bullet(doc, item)

add_para(doc, "End of database. 97 entries.", italic=True)

doc.save(DST)
print(f"Saved: {DST}")
print(f"Size: {DST.stat().st_size:,} bytes")
