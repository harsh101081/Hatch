from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_heading_styled(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_bold_paragraph(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def add_normal_paragraph(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    return p


def add_page_break():
    doc.add_page_break()



# ========================
# COVER PAGE
# ========================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('INTERNSHIP TRAINING REPORT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Financial Analyst Intern')
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()
company = doc.add_paragraph()
company.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = company.add_run('FUNDSWEB')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
dept = doc.add_paragraph()
dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dept.add_run('Department: Finance, Investment Research & Analytics')
run.font.size = Pt(14)

doc.add_paragraph()
duration = doc.add_paragraph()
duration.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = duration.add_run('Duration: 8 Days Training Program')
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Intern details placeholder
details_heading = doc.add_paragraph()
details_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details_heading.add_run('INTERN DETAILS')
run.bold = True
run.font.size = Pt(14)

details = [
    'Name: ___________________________',
    'Roll No.: ___________________________',
    'College/University: ___________________________',
    'Course: ___________________________',
    'Semester/Year: ___________________________',
    'Contact No.: ___________________________',
    'Email: ___________________________',
    'Reporting Manager: ___________________________',
]

for d in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(d)
    run.font.size = Pt(12)

add_page_break()


# ========================
# TABLE OF CONTENTS
# ========================
add_heading_styled('TABLE OF CONTENTS', level=1)
doc.add_paragraph()

toc_items = [
    '1. Day 1: General Orientation',
    '2. Day 2: Portfolio Management (Digital Resume)',
    '3. Day 3: Strategic Investment Planning',
    '4. Day 4: Stock Analysis',
    '5. Day 5: Fundamental Analysis',
    '6. Day 6: Technical Analysis',
    '7. Day 7: Calculators Creation and Case Study Solving',
    '8. Day 8: Business Report Creation',
    '9. Overall Internship Experience',
    '10. Skills Developed',
    '11. Challenges Faced & Solutions Implemented',
    '12. Practical Exposure Gained',
    '13. Conclusion',
    '14. Acknowledgement',
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1)

add_page_break()


# ========================
# DAY 1: GENERAL ORIENTATION
# ========================
add_heading_styled('DAY 1: GENERAL ORIENTATION', level=1)

add_bold_paragraph('1.1 Introduction')
add_normal_paragraph(
    'The first day of the internship at Fundsweb marked the beginning of a structured '
    '8-day training programme designed to provide hands-on exposure to the finance and '
    'investment research industry. The orientation session served as a foundation for '
    'understanding the company\'s operations, its service offerings, and the broader '
    'landscape of financial services in India.'
)

add_bold_paragraph('1.2 Aim')
add_normal_paragraph(
    'To gain a comprehensive understanding of Fundsweb as an organisation, its core '
    'business operations, the structure of its various departments, and the expectations '
    'set for the internship training programme.'
)

add_bold_paragraph('1.3 Objectives')
add_bullet('To understand the history, vision, and mission of Fundsweb.')
add_bullet('To learn about the organisational hierarchy and departmental functions.')
add_bullet('To gain awareness of the financial products and services offered by the company.')
add_bullet('To understand the role of a Financial Analyst Intern within the organisation.')
add_bullet('To get acquainted with the tools, platforms, and resources used in daily operations.')
add_bullet('To set clear expectations and deliverables for the training period.')

add_bold_paragraph('1.4 Tasks Performed')
add_normal_paragraph(
    'The day commenced with a formal introduction session conducted by the team lead, '
    'where we were briefed about the company\'s journey since its inception. Fundsweb '
    'operates as a financial services and investment research firm that provides advisory, '
    'portfolio management, and wealth planning solutions to retail and institutional clients. '
    'We were given a walkthrough of the company\'s website and its digital platforms to '
    'understand the client interface and service delivery model.'
)
add_normal_paragraph(
    'The session then moved to a detailed explanation of the organisational structure. '
    'We learnt about the different departments including research and analytics, client '
    'relations, compliance, and technology. Each department\'s contribution to the overall '
    'functioning was explained through real examples. We were also introduced to the '
    'compliance and regulatory framework that governs financial services in India, '
    'including SEBI guidelines and AMFI regulations.'
)
add_normal_paragraph(
    'Post lunch, we had an interactive Q&A session where doubts regarding the training '
    'schedule, evaluation criteria, and expected outcomes were clarified. We were also '
    'assigned login credentials for internal tools and given reading material on the '
    'basics of financial markets to prepare for the upcoming sessions.'
)

add_bold_paragraph('1.5 Key Learnings')
add_bullet('Fundsweb operates across multiple verticals including mutual funds, equity research, and insurance advisory.')
add_bullet('Regulatory compliance is a non-negotiable aspect of financial services delivery.')
add_bullet('Understanding organisational structure helps in appreciating how different functions contribute to a common goal.')
add_bullet('The importance of professional conduct and ethical standards in the finance industry.')
add_bullet('Familiarisation with industry terminology used in day-to-day operations.')

add_bold_paragraph('1.6 Findings & Conclusion')
add_normal_paragraph(
    'The orientation day was instrumental in setting the right context for the entire '
    'training programme. It became evident that Fundsweb places significant emphasis on '
    'research-driven investment advice and client-centric service delivery. The structured '
    'approach to the training, with clearly defined daily themes, reflected the company\'s '
    'commitment to developing well-rounded financial professionals. The session helped me '
    'understand that the finance industry demands not just technical knowledge but also '
    'strong communication skills, ethical awareness, and the ability to adapt to rapidly '
    'changing market conditions. Overall, the first day provided a solid foundation upon '
    'which subsequent learning was built.'
)

add_page_break()


# ========================
# DAY 2: PORTFOLIO MANAGEMENT (DIGITAL RESUME)
# ========================
add_heading_styled('DAY 2: PORTFOLIO MANAGEMENT (DIGITAL RESUME)', level=1)

add_bold_paragraph('2.1 Introduction')
add_normal_paragraph(
    'The second day focused on two interconnected themes — portfolio management principles '
    'and the creation of a digital resume. The rationale behind combining these two topics '
    'was to draw a parallel between managing a financial portfolio and managing one\'s own '
    'professional profile. Just as diversification reduces risk in investments, a '
    'well-rounded professional profile increases career opportunities.'
)

add_bold_paragraph('2.2 Aim')
add_normal_paragraph(
    'To understand the fundamental principles of portfolio management including asset '
    'allocation, diversification, and risk-return trade-offs, and to apply similar '
    'principles in constructing a compelling digital resume.'
)

add_bold_paragraph('2.3 Objectives')
add_bullet('To understand the concept of a portfolio and its components.')
add_bullet('To learn about asset allocation strategies across different asset classes.')
add_bullet('To comprehend the role of diversification in managing investment risk.')
add_bullet('To analyse the risk-return relationship in portfolio construction.')
add_bullet('To create a professional digital resume highlighting relevant skills and experiences.')
add_bullet('To understand personal branding in the context of financial services careers.')

add_bold_paragraph('2.4 Tasks Performed')
add_normal_paragraph(
    'The morning session began with a theoretical overview of portfolio management. '
    'We studied how institutional and retail investors construct portfolios based on '
    'their risk appetite, investment horizon, and financial goals. The trainer explained '
    'the concept of asset classes — equity, debt, gold, and real estate — and how the '
    'proportion allocated to each depends on the investor\'s profile. We worked through '
    'examples of conservative, moderate, and aggressive portfolios to understand how '
    'allocation shifts based on risk tolerance.'
)
add_normal_paragraph(
    'Diversification was discussed in depth with practical examples. We analysed how '
    'holding stocks across different sectors such as IT, banking, FMCG, and pharma '
    'reduces the impact of sector-specific downturns on the overall portfolio. The '
    'concept of correlation between assets was introduced, explaining why combining '
    'negatively correlated assets provides better risk-adjusted returns over time.'
)
add_normal_paragraph(
    'In the afternoon, we pivoted to building our digital resumes. Using online '
    'platforms and templates, we created professional profiles that highlighted our '
    'academic background, skills, certifications, and internship experiences. The '
    'trainer emphasised that a digital presence is today\'s equivalent of a portfolio — '
    'it represents what you bring to the table. We received feedback on our drafts and '
    'made revisions to improve clarity and impact.'
)

add_bold_paragraph('2.5 Key Learnings')
add_bullet('A diversified portfolio significantly reduces unsystematic risk without proportionally reducing expected returns.')
add_bullet('Asset allocation is the single most important decision in portfolio construction.')
add_bullet('The risk-return trade-off means higher expected returns always come with higher volatility.')
add_bullet('A digital resume must be concise, relevant, and tailored to the target industry.')
add_bullet('Professional branding is as important in finance careers as technical competence.')

add_bold_paragraph('2.6 Findings & Conclusion')
add_normal_paragraph(
    'The session effectively demonstrated that portfolio management is both a science '
    'and an art. The quantitative aspect involves calculating expected returns, standard '
    'deviations, and correlations, while the qualitative aspect involves understanding '
    'investor behaviour and market sentiment. The parallel drawn with digital resume '
    'building was insightful — just as one diversifies a portfolio to manage risk, one '
    'should diversify skills and experiences to remain competitive in the job market. '
    'The practical exercise of creating a digital resume reinforced the importance of '
    'presenting oneself professionally in an increasingly digital world. This day\'s '
    'learning laid the groundwork for understanding how investments should be structured '
    'and how personal career development mirrors financial planning principles.'
)

add_page_break()


# ========================
# DAY 3: STRATEGIC INVESTMENT PLANNING
# ========================
add_heading_styled('DAY 3: STRATEGIC INVESTMENT PLANNING', level=1)

add_bold_paragraph('3.1 Introduction')
add_normal_paragraph(
    'Day three was dedicated to strategic investment planning — the process of aligning '
    'one\'s financial resources with long-term goals through disciplined and informed '
    'decision-making. The session bridged the gap between theoretical portfolio concepts '
    'and their practical application in real-world financial planning scenarios.'
)

add_bold_paragraph('3.2 Aim')
add_normal_paragraph(
    'To develop a practical understanding of investment planning by learning how to '
    'set financial goals, assess risk profiles, determine investment horizons, and '
    'construct goal-based investment strategies.'
)

add_bold_paragraph('3.3 Objectives')
add_bullet('To understand the process of setting SMART financial goals.')
add_bullet('To learn the methodology of risk profiling for individual investors.')
add_bullet('To comprehend the relationship between investment horizon and asset selection.')
add_bullet('To study wealth creation strategies including systematic investment plans (SIPs).')
add_bullet('To develop a goal-based investment plan for a hypothetical client.')
add_bullet('To understand the impact of inflation on long-term financial planning.')

add_bold_paragraph('3.4 Tasks Performed')
add_normal_paragraph(
    'The session started with a discussion on why most individuals fail at wealth creation — '
    'primarily due to the absence of clearly defined financial goals. We studied the SMART '
    'framework (Specific, Measurable, Achievable, Relevant, Time-bound) and applied it to '
    'common financial objectives such as buying a house, children\'s education, and '
    'retirement planning. Each goal was broken down into its time horizon, required corpus, '
    'and the monthly investment needed to achieve it.'
)
add_normal_paragraph(
    'Risk profiling was covered next. We filled out sample risk assessment questionnaires '
    'and classified investors into categories — conservative, moderate, and aggressive. '
    'The trainer explained how risk profiling directly influences asset allocation decisions. '
    'For instance, a young professional with a 20-year horizon can afford higher equity '
    'exposure compared to someone nearing retirement who needs capital preservation.'
)
add_normal_paragraph(
    'We then worked on creating complete financial plans for hypothetical clients. This '
    'involved calculating future values of goals adjusted for inflation, determining the '
    'required rate of return, selecting appropriate investment vehicles (mutual funds, '
    'PPF, NPS, equities), and mapping each investment to a specific goal. The exercise '
    'gave us a hands-on feel of what financial advisors do when onboarding new clients.'
)

add_bold_paragraph('3.5 Key Learnings')
add_bullet('Goal-based investing is more effective than ad-hoc investment decisions.')
add_bullet('Inflation erodes purchasing power significantly over long periods, making growth-oriented investments essential.')
add_bullet('Risk profiling ensures that investment recommendations are suitable for the client\'s temperament and capacity.')
add_bullet('Systematic Investment Plans (SIPs) leverage rupee cost averaging to reduce timing risk.')
add_bullet('Investment planning is a continuous process that requires periodic review and rebalancing.')

add_bold_paragraph('3.6 Findings & Conclusion')
add_normal_paragraph(
    'Strategic investment planning is the backbone of sound financial management. The '
    'session made it clear that without a structured plan, even high-income individuals '
    'can fail to achieve financial independence. The exercise of building a financial plan '
    'from scratch — starting with goal identification, moving through risk assessment, '
    'and ending with product selection — provided an end-to-end understanding of the '
    'advisory process. I found that the most challenging part is not the calculation '
    'itself but understanding the client\'s real needs and translating them into '
    'actionable investment decisions. This day reinforced that financial planning is '
    'as much about behavioural understanding as it is about numbers and projections.'
)

add_page_break()


# ========================
# DAY 4: STOCK ANALYSIS
# ========================
add_heading_styled('DAY 4: STOCK ANALYSIS', level=1)

add_bold_paragraph('4.1 Introduction')
add_normal_paragraph(
    'The fourth day introduced the practice of stock analysis — the systematic evaluation '
    'of individual stocks to determine their suitability for investment. This session '
    'focused on the qualitative and quantitative parameters that analysts use to screen, '
    'shortlist, and evaluate companies before recommending them to investors.'
)

add_bold_paragraph('4.2 Aim')
add_normal_paragraph(
    'To learn the methodologies and frameworks used in stock analysis, including sector '
    'analysis, company evaluation, competitive positioning, and financial performance '
    'assessment, to make informed investment decisions.'
)

add_bold_paragraph('4.3 Objectives')
add_bullet('To understand the stock screening process and the criteria used for filtering stocks.')
add_bullet('To learn sector analysis and identify sectors with growth potential.')
add_bullet('To evaluate companies based on market capitalisation, revenue growth, and profitability.')
add_bullet('To understand the concept of economic moats and competitive advantages.')
add_bullet('To apply stock analysis frameworks to real companies listed on Indian stock exchanges.')
add_bullet('To differentiate between large-cap, mid-cap, and small-cap stocks and their risk profiles.')

add_bold_paragraph('4.4 Tasks Performed')
add_normal_paragraph(
    'The morning session began with stock screening. We used online screeners to filter '
    'stocks based on parameters such as market capitalisation, P/E ratio, revenue growth, '
    'promoter holding, and debt levels. The trainer demonstrated how to set screening '
    'criteria based on investment style — value investing requires low P/E and P/B ratios, '
    'while growth investing focuses on high revenue and earnings growth rates.'
)
add_normal_paragraph(
    'We then moved to sector analysis. The Indian market was divided into major sectors — '
    'banking, IT, pharmaceuticals, FMCG, automobiles, and energy. We studied the current '
    'macroeconomic factors affecting each sector, regulatory changes, and sectoral trends. '
    'For example, we discussed how rising interest rates benefit banking stocks while '
    'negatively impacting real estate and automobile sectors due to increased borrowing costs.'
)
add_normal_paragraph(
    'In the afternoon, we picked two companies from different sectors and performed a '
    'detailed comparative analysis. We examined their revenue trends over five years, '
    'profit margins, return ratios, debt levels, and management quality. We also assessed '
    'their competitive advantages — brand strength, distribution network, patent portfolios, '
    'or cost leadership — to determine which company had a more sustainable business model. '
    'The findings were compiled into a brief stock research note format.'
)

add_bold_paragraph('4.5 Key Learnings')
add_bullet('Stock screening is the first step in narrowing down the investment universe to manageable candidates.')
add_bullet('Sector analysis provides context — a good company in a declining sector may still underperform.')
add_bullet('Market capitalisation determines the risk-reward profile: small-caps offer higher growth but with greater volatility.')
add_bullet('Competitive advantage or economic moat is critical for long-term wealth creation through equities.')
add_bullet('Management quality and corporate governance are qualitative factors that significantly impact stock performance.')

add_bold_paragraph('4.6 Findings & Conclusion')
add_normal_paragraph(
    'Stock analysis requires a blend of quantitative rigour and qualitative judgement. '
    'The day\'s exercises demonstrated that no single metric can determine whether a stock '
    'is a good investment — it is the combination of financial health, growth prospects, '
    'competitive positioning, and fair valuation that makes a compelling case. I observed '
    'that the process is iterative; initial screening narrows the universe, detailed '
    'analysis builds conviction, and continuous monitoring ensures the investment thesis '
    'remains valid. The practical exposure to real stock data and screeners was particularly '
    'valuable as it connected classroom concepts to market realities. This session built '
    'a strong foundation for the more detailed fundamental and technical analysis sessions '
    'that followed.'
)

add_page_break()


# ========================
# DAY 5: FUNDAMENTAL ANALYSIS
# ========================
add_heading_styled('DAY 5: FUNDAMENTAL ANALYSIS', level=1)

add_bold_paragraph('5.1 Introduction')
add_normal_paragraph(
    'Day five delved deep into fundamental analysis — the cornerstone of long-term '
    'investment decision-making. While Day 4 provided a broad overview of stock analysis, '
    'this session focused specifically on reading financial statements, calculating key '
    'financial ratios, and using valuation techniques to determine whether a stock is '
    'undervalued, fairly valued, or overvalued.'
)

add_bold_paragraph('5.2 Aim')
add_normal_paragraph(
    'To develop proficiency in fundamental analysis by learning to interpret financial '
    'statements, compute and analyse critical financial ratios, and apply valuation '
    'methodologies to identify fundamentally strong companies for investment.'
)

add_bold_paragraph('5.3 Objectives')
add_bullet('To understand the three core financial statements — Balance Sheet, Income Statement, and Cash Flow Statement.')
add_bullet('To calculate and interpret key ratios: P/E, P/B, ROE, ROCE, Debt-to-Equity, and EPS.')
add_bullet('To learn valuation techniques including DCF and relative valuation.')
add_bullet('To identify red flags in financial statements that indicate poor financial health.')
add_bullet('To apply fundamental analysis to evaluate at least two listed companies.')
add_bullet('To understand how fundamental analysis differs from and complements technical analysis.')

add_bold_paragraph('5.4 Tasks Performed')
add_normal_paragraph(
    'The session began with a detailed walkthrough of financial statements. We took '
    'the annual report of a well-known listed company and went through each component '
    'of the balance sheet — assets (current and non-current), liabilities, and shareholders\' '
    'equity. The income statement was analysed for revenue trends, operating margins, '
    'depreciation, interest costs, and net profit. The cash flow statement was studied '
    'to understand operating cash flows, capital expenditure, and free cash flow generation.'
)
add_normal_paragraph(
    'Next, we calculated financial ratios. The Price-to-Earnings (P/E) ratio helped us '
    'gauge how much the market is willing to pay per rupee of earnings. The Price-to-Book '
    '(P/B) ratio indicated whether the stock was trading above or below its book value. '
    'Return on Equity (ROE) and Return on Capital Employed (ROCE) revealed how efficiently '
    'the company was using shareholder funds and total capital respectively. The '
    'Debt-to-Equity ratio flagged companies with excessive leverage, while Earnings '
    'Per Share (EPS) provided a per-share profitability measure.'
)
add_normal_paragraph(
    'In the practical exercise, we compared two companies from the same sector using '
    'all the above ratios and concluded which one appeared more fundamentally sound. '
    'We also discussed the Discounted Cash Flow (DCF) model briefly, understanding how '
    'future cash flows are discounted back to present value to arrive at an intrinsic '
    'value estimate. The session concluded with a discussion on the limitations of '
    'fundamental analysis, including the time lag in financial reporting and the '
    'assumptions inherent in valuation models.'
)

add_bold_paragraph('5.5 Key Learnings')
add_bullet('The balance sheet shows financial position at a point in time, while the income statement shows performance over a period.')
add_bullet('Cash flow from operations is often a more reliable indicator of financial health than net profit.')
add_bullet('A high ROE combined with low debt suggests efficient capital utilisation.')
add_bullet('P/E ratios must be compared within the same sector for meaningful interpretation.')
add_bullet('Fundamental analysis helps identify companies with strong financials that are likely to generate long-term wealth.')
add_bullet('No single ratio tells the complete story — ratios must be analysed collectively.')

add_bold_paragraph('5.6 Findings & Conclusion')
add_normal_paragraph(
    'Fundamental analysis emerged as a powerful tool for identifying quality companies '
    'at reasonable valuations. The hands-on exercise of reading actual financial statements '
    'and computing ratios brought theoretical concepts to life. I found that while ratio '
    'analysis provides quick insights, understanding the narrative behind the numbers — '
    'why margins expanded, what drove revenue growth, how the company managed its working '
    'capital — is equally important. The session reinforced that successful long-term '
    'investing is about buying good businesses at fair prices, and fundamental analysis '
    'provides the framework to make that assessment. The combination of quantitative '
    'ratios and qualitative business understanding forms the bedrock of value investing '
    'philosophy that many successful investors follow.'
)

add_page_break()


# ========================
# DAY 6: TECHNICAL ANALYSIS
# ========================
add_heading_styled('DAY 6: TECHNICAL ANALYSIS', level=1)

add_bold_paragraph('6.1 Introduction')
add_normal_paragraph(
    'The sixth day introduced technical analysis — a methodology that uses historical '
    'price and volume data to forecast future price movements. Unlike fundamental analysis '
    'which focuses on intrinsic value, technical analysis is concerned with market '
    'psychology, trends, and patterns. This session equipped us with the tools and '
    'indicators used by traders and short-term investors to time their entries and exits.'
)

add_bold_paragraph('6.2 Aim')
add_normal_paragraph(
    'To understand the principles and tools of technical analysis, including chart '
    'patterns, indicators, and trading strategies, and to apply them for identifying '
    'market trends and potential trading opportunities.'
)

add_bold_paragraph('6.3 Objectives')
add_bullet('To learn to read and interpret candlestick charts and their patterns.')
add_bullet('To understand support and resistance levels and their significance in trading.')
add_bullet('To study trend analysis and identify bullish, bearish, and sideways trends.')
add_bullet('To compute and interpret Moving Averages (SMA and EMA) for trend confirmation.')
add_bullet('To apply RSI and MACD indicators for momentum and signal analysis.')
add_bullet('To understand volume analysis as confirmation of price movements.')

add_bold_paragraph('6.4 Tasks Performed')
add_normal_paragraph(
    'The morning session began with candlestick charts. We studied individual candlestick '
    'patterns — doji, hammer, engulfing, and shooting star — and understood what each '
    'pattern signals about market sentiment. The trainer used live charts on TradingView '
    'to demonstrate how these patterns appear at key support and resistance zones and '
    'often precede trend reversals.'
)
add_normal_paragraph(
    'Support and resistance were explained as psychological price levels where buying '
    'or selling pressure historically concentrates. We practised drawing horizontal '
    'support and resistance lines on charts of Nifty 50 and select large-cap stocks. '
    'Trend lines were drawn to identify the direction of the prevailing trend — '
    'connecting higher lows in an uptrend and lower highs in a downtrend.'
)
add_normal_paragraph(
    'In the second half, we covered technical indicators. The 50-day and 200-day Simple '
    'Moving Averages (SMA) were used to identify long-term trends and golden/death cross '
    'signals. The Relative Strength Index (RSI) was explained as a momentum oscillator '
    'that identifies overbought (above 70) and oversold (below 30) conditions. The MACD '
    '(Moving Average Convergence Divergence) was studied for signal line crossovers and '
    'histogram divergences. We practised applying these indicators to real charts and '
    'documenting potential entry and exit points based on multiple indicator confluence.'
)

add_bold_paragraph('6.5 Key Learnings')
add_bullet('Technical analysis is based on the premise that price discounts everything and history tends to repeat.')
add_bullet('Support and resistance levels serve as crucial decision-making zones for traders.')
add_bullet('Moving averages smooth out price noise and help identify the underlying trend direction.')
add_bullet('RSI and MACD are complementary indicators — RSI measures momentum while MACD confirms trend changes.')
add_bullet('Volume confirmation is essential; price movements on high volume are more reliable.')
add_bullet('No single indicator is foolproof; combining multiple tools increases the probability of successful trades.')

add_bold_paragraph('6.6 Findings & Conclusion')
add_normal_paragraph(
    'Technical analysis provides a structured approach to understanding market behaviour '
    'through price action. The session demonstrated that while fundamental analysis answers '
    '"what to buy," technical analysis largely answers "when to buy." The practical '
    'exercise of identifying patterns and applying indicators on live charts was extremely '
    'engaging and provided clarity on how professional traders make decisions. I found '
    'that discipline and risk management are more important than the accuracy of any '
    'individual signal. The combination of trend identification through moving averages '
    'and momentum confirmation through RSI/MACD creates a robust framework for timing '
    'investment decisions. This session complemented the fundamental analysis training '
    'by adding the dimension of market timing to our analytical toolkit.'
)

add_page_break()


# ========================
# DAY 7: CALCULATORS CREATION AND CASE STUDY SOLVING
# ========================
add_heading_styled('DAY 7: CALCULATORS CREATION AND CASE STUDY SOLVING', level=1)

add_bold_paragraph('7.1 Introduction')
add_normal_paragraph(
    'Day seven was the most hands-on session of the training programme. We were tasked '
    'with building financial calculators and solving a comprehensive case study. This '
    'day tested our ability to apply the theoretical knowledge gained over the previous '
    'six days to practical problem-solving scenarios that financial analysts encounter '
    'in their daily work.'
)

add_bold_paragraph('7.2 Aim')
add_normal_paragraph(
    'To design and build functional financial calculators for common investment planning '
    'tasks, and to solve a real-world case study involving personal financial planning '
    'and investment decision-making.'
)

add_bold_paragraph('7.3 Objectives')
add_bullet('To build a SIP Calculator that computes future value of systematic investments.')
add_bullet('To create a SWP Calculator for planning systematic withdrawal strategies.')
add_bullet('To develop a Financial Ratio Calculator for quick ratio computation.')
add_bullet('To design a Position Size and Risk Management Calculator for trading decisions.')
add_bullet('To solve a case study integrating financial assessment, risk analysis, and investment strategy.')
add_bullet('To present findings in a structured and professional format.')

add_bold_paragraph('7.4 Tasks Performed')

add_bold_paragraph('7.4.1 SIP Calculator')
add_normal_paragraph(
    'The SIP (Systematic Investment Plan) Calculator was designed to help investors '
    'estimate the future value of their regular monthly investments. The inputs included '
    'monthly investment amount, expected annual rate of return, and investment tenure in '
    'years. The calculator used the compound interest formula adapted for periodic '
    'contributions: FV = P x [((1+r)^n - 1) / r] x (1+r), where P is the monthly '
    'contribution, r is the monthly rate of return, and n is the total number of months. '
    'The output displayed total amount invested, estimated returns, and total corpus at '
    'maturity. This tool is practically useful for clients who want to visualise how '
    'small monthly savings compound into significant wealth over time.'
)

add_bold_paragraph('7.4.2 SWP Calculator')
add_normal_paragraph(
    'The SWP (Systematic Withdrawal Plan) Calculator was built for retirees or investors '
    'who want regular income from their accumulated corpus. Inputs included the initial '
    'investment amount, expected annual return on the remaining corpus, monthly withdrawal '
    'amount, and withdrawal period. The calculator computed the remaining corpus after '
    'each year, total amount withdrawn, and whether the corpus would sustain the planned '
    'withdrawals for the desired period. This tool helps in retirement planning by showing '
    'whether a given corpus can support a particular lifestyle without premature depletion.'
)

add_bold_paragraph('7.4.3 Financial Ratio Calculator')
add_normal_paragraph(
    'This calculator automated the computation of key financial ratios from raw financial '
    'data inputs. Users could input values such as net profit, total equity, total assets, '
    'current liabilities, market price, and earnings per share. The calculator then '
    'computed ratios including P/E, P/B, ROE, ROCE, Debt-to-Equity, current ratio, and '
    'net profit margin. This tool saves considerable time during stock analysis by '
    'eliminating manual calculations and reducing the probability of arithmetic errors.'
)

add_bold_paragraph('7.4.4 Position Size & Risk Management Calculator')
add_normal_paragraph(
    'This calculator was designed for traders to determine the appropriate position size '
    'based on their risk tolerance. Inputs included total trading capital, risk percentage '
    'per trade (typically 1-2%), entry price, and stop-loss price. The calculator computed '
    'the maximum number of shares to buy, total capital at risk, and the risk-reward ratio '
    'based on the target price. This tool enforces disciplined trading by ensuring that '
    'no single trade can cause disproportionate damage to the trading account.'
)


add_bold_paragraph('7.5 Case Study: Personal Financial Planning for a Young Professional')

add_bold_paragraph('Problem Statement')
add_normal_paragraph(
    'Mr. Rahul Mehta, aged 28, is a software engineer earning a monthly salary of '
    'Rs. 85,000. He has no existing investments other than his EPF contributions. He '
    'has the following financial goals: (1) Build an emergency fund of 6 months\' expenses '
    'within one year, (2) Accumulate Rs. 25 lakhs for a home down payment in 5 years, '
    '(3) Create a retirement corpus of Rs. 5 crores by age 55. His monthly expenses are '
    'Rs. 40,000, and he has an education loan EMI of Rs. 12,000 with 2 years remaining. '
    'He has a moderate risk appetite and is looking for a comprehensive investment plan.'
)

add_bold_paragraph('Financial Assessment')
add_normal_paragraph(
    'Monthly surplus available for investment: Rs. 85,000 - Rs. 40,000 (expenses) - '
    'Rs. 12,000 (EMI) = Rs. 33,000. After the loan closure in 2 years, available surplus '
    'increases to Rs. 45,000. Current net worth is negligible apart from EPF accumulation '
    'of approximately Rs. 2.5 lakhs. No existing insurance coverage beyond employer-provided '
    'group insurance.'
)

add_bold_paragraph('Risk Identification')
add_bullet('No emergency fund exposes him to liquidity risk in case of job loss or medical emergency.')
add_bullet('Lack of term insurance creates financial vulnerability for dependents.')
add_bullet('Single source of income with no passive income streams.')
add_bullet('Education loan obligation reduces current investment capacity.')
add_bullet('Inflation risk on the home purchase goal given rising real estate prices.')

add_bold_paragraph('Investment Strategy')
add_normal_paragraph(
    'Phase 1 (Months 1-12): Priority allocation to emergency fund. Invest Rs. 20,000/month '
    'in a liquid fund to build Rs. 2.4 lakh emergency corpus. Simultaneously start SIPs '
    'of Rs. 10,000 in a large-cap equity fund for the home purchase goal. Purchase term '
    'insurance of Rs. 1 crore (premium approximately Rs. 800/month) and health insurance '
    'of Rs. 10 lakhs.'
)
add_normal_paragraph(
    'Phase 2 (Months 13-24): Emergency fund achieved. Redirect Rs. 20,000 to the home '
    'purchase goal through a balanced advantage fund (total monthly investment for this '
    'goal becomes Rs. 30,000). Continue existing SIPs.'
)
add_normal_paragraph(
    'Phase 3 (Month 25 onwards): Loan EMI of Rs. 12,000 freed up. Total investible surplus '
    'becomes Rs. 45,000. Allocate Rs. 30,000 towards home goal, Rs. 15,000 towards '
    'retirement via equity mutual funds (flexi-cap and mid-cap categories for long-term '
    'compounding).'
)

add_bold_paragraph('Solution Provided')
add_normal_paragraph(
    'A phased investment plan was recommended that prioritised immediate needs (emergency fund, '
    'insurance) before long-term goals. The home purchase goal was mapped to balanced '
    'advantage funds offering moderate risk over 5 years. The retirement goal, with a '
    '27-year horizon, was allocated to higher equity exposure for maximum compounding benefit.'
)

add_bold_paragraph('Outcome')
add_normal_paragraph(
    'Based on projected returns of 10% for equity and 7% for balanced funds: Emergency fund '
    'target of Rs. 2.4 lakhs achievable within 12 months. Home purchase corpus of Rs. 25+ '
    'lakhs achievable in 5 years with Rs. 30,000 monthly SIP at 10% CAGR. Retirement corpus '
    'projection of Rs. 5.2 crores achievable by age 55 through Rs. 15,000 monthly SIP '
    'increasing by 10% annually.'
)

add_bold_paragraph('Conclusion')
add_normal_paragraph(
    'The case study demonstrated the importance of structured financial planning. By '
    'prioritising goals, assessing risk, and selecting appropriate products, even a '
    'moderate income can generate substantial long-term wealth. The key takeaway was that '
    'starting early and investing consistently matters more than the absolute amount invested.'
)

add_bold_paragraph('7.6 Key Learnings')
add_bullet('Financial calculators simplify complex computations and improve advisory efficiency.')
add_bullet('Position sizing is critical — it determines survival in trading more than entry accuracy.')
add_bullet('Real-world financial planning requires balancing multiple competing goals with limited resources.')
add_bullet('Insurance is a prerequisite to investment planning, not an afterthought.')
add_bullet('Phased investment strategies adapt to changing financial circumstances over time.')

add_bold_paragraph('7.7 Findings & Conclusion')
add_normal_paragraph(
    'This was arguably the most practical day of the entire training. Building calculators '
    'from scratch reinforced the mathematical foundations underlying investment planning, '
    'while the case study integrated all concepts — goal setting, risk assessment, product '
    'selection, and implementation — into a cohesive solution. The experience of solving '
    'a realistic client scenario gave me confidence that the knowledge gained during this '
    'internship has direct professional applicability. Financial planning is not about '
    'complex strategies but about disciplined execution of a well-thought-out plan tailored '
    'to individual circumstances.'
)

add_page_break()


# ========================
# DAY 8: BUSINESS REPORT CREATION
# ========================
add_heading_styled('DAY 8: BUSINESS REPORT CREATION', level=1)

add_bold_paragraph('8.1 Introduction')
add_normal_paragraph(
    'The final day of the training programme was dedicated to business report writing — '
    'an essential skill for any finance professional. The ability to compile research '
    'findings, data analysis, and recommendations into a clear, structured, and '
    'professional document is what distinguishes a competent analyst from an exceptional '
    'one. This session covered the principles of report writing, data presentation, and '
    'professional documentation standards.'
)

add_bold_paragraph('8.2 Aim')
add_normal_paragraph(
    'To learn the structure, methodology, and best practices of business report writing '
    'in the context of finance and investment research, and to develop the ability to '
    'present complex financial information in a clear and actionable format.'
)

add_bold_paragraph('8.3 Objectives')
add_bullet('To understand the standard structure of a business/financial report.')
add_bullet('To learn research methodology including primary and secondary data collection.')
add_bullet('To develop skills in data interpretation and drawing meaningful conclusions.')
add_bullet('To practise presenting quantitative findings alongside qualitative insights.')
add_bullet('To understand the importance of visual aids — charts, tables, and graphs — in reports.')
add_bullet('To compile the internship experience into a comprehensive training report.')

add_bold_paragraph('8.4 Tasks Performed')
add_normal_paragraph(
    'The morning session covered the theoretical framework of business report writing. '
    'We studied the standard components of a financial report: executive summary, '
    'introduction, methodology, findings, analysis, recommendations, and conclusion. '
    'The trainer emphasised that a good report must be reader-centric — it should '
    'anticipate the reader\'s questions and answer them in a logical sequence.'
)
add_normal_paragraph(
    'We discussed research methodology in the context of financial research. Primary '
    'data sources include company filings, annual reports, management commentary, and '
    'investor presentations. Secondary sources include industry reports, broker research, '
    'news articles, and regulatory filings. The importance of cross-referencing data from '
    'multiple sources to ensure accuracy was stressed. We also covered data interpretation '
    '— how to identify trends, outliers, and patterns in financial data and translate '
    'them into meaningful narratives.'
)
add_normal_paragraph(
    'In the practical session, we began compiling our internship training report. This '
    'involved organising the notes and exercises from all eight days into a structured '
    'document with proper headings, numbered sections, and consistent formatting. We '
    'were guided on how to write findings and conclusions that add value rather than '
    'merely restating what was done. The session also covered professional documentation '
    'standards including citation, referencing, and maintaining objectivity in reporting.'
)

add_bold_paragraph('8.5 Key Learnings')
add_bullet('A well-structured report guides the reader through complex information without confusion.')
add_bullet('Executive summaries should be written last but placed first — they capture the essence of the entire report.')
add_bullet('Data without interpretation is just numbers; the analyst\'s job is to provide context and meaning.')
add_bullet('Visual representation of data (charts, graphs) significantly improves comprehension and retention.')
add_bullet('Professional reports must maintain objectivity — presenting facts and analysis rather than opinions.')

add_bold_paragraph('8.6 Findings & Conclusion')
add_normal_paragraph(
    'Business report writing is a skill that ties together all aspects of financial '
    'analysis. The session made it evident that even the most brilliant analysis is '
    'worthless if it cannot be communicated effectively. In the finance industry, reports '
    'drive decisions — investment committees approve recommendations based on research '
    'reports, clients make financial decisions based on advisory documents, and regulators '
    'assess compliance through filed reports. The ability to write clearly, present data '
    'compellingly, and arrive at well-reasoned conclusions is therefore not a peripheral '
    'skill but a core competency. This final day brought closure to the training by '
    'giving us the tools to document and communicate everything we had learnt throughout '
    'the programme.'
)

add_page_break()


# ========================
# FINAL SECTIONS
# ========================
add_heading_styled('OVERALL INTERNSHIP EXPERIENCE', level=1)
add_normal_paragraph(
    'The 8-day training programme at Fundsweb was an enriching and transformative experience '
    'that provided comprehensive exposure to the finance and investment research industry. '
    'The programme was meticulously designed to cover the full spectrum of financial analysis — '
    'from understanding organisational operations on Day 1 to producing professional business '
    'reports on Day 8. Each day built upon the previous one, creating a logical progression '
    'of knowledge and skills.'
)
add_normal_paragraph(
    'What stood out most about the programme was its emphasis on practical application. '
    'Rather than limiting the training to theoretical lectures, the facilitators ensured '
    'that every concept was reinforced through hands-on exercises, real market data, and '
    'industry-relevant case studies. The exposure to actual stock screeners, charting '
    'platforms, and financial calculators bridged the gap between academic learning and '
    'professional practice.'
)
add_normal_paragraph(
    'The internship also provided valuable insights into the soft skills required in '
    'finance — professional communication, analytical thinking, attention to detail, and '
    'the ability to synthesise large amounts of information into actionable recommendations. '
    'Working in a professional environment, meeting deadlines, and collaborating with '
    'fellow interns enhanced my readiness for a full-time role in the industry.'
)

add_page_break()

# SKILLS DEVELOPED
add_heading_styled('SKILLS DEVELOPED', level=1)
add_bold_paragraph('Technical Skills')
add_bullet('Financial statement analysis and ratio computation')
add_bullet('Stock screening and sector analysis')
add_bullet('Technical chart reading and indicator application')
add_bullet('Portfolio construction and asset allocation')
add_bullet('Financial calculator development')
add_bullet('Investment planning and goal mapping')
add_bullet('Business report writing and data presentation')

add_bold_paragraph('Analytical Skills')
add_bullet('Critical evaluation of investment opportunities')
add_bullet('Risk assessment and management')
add_bullet('Data interpretation and trend identification')
add_bullet('Comparative analysis of financial instruments')
add_bullet('Case study problem-solving methodology')

add_bold_paragraph('Professional Skills')
add_bullet('Professional communication and presentation')
add_bullet('Time management and deadline adherence')
add_bullet('Digital proficiency with financial tools and platforms')
add_bullet('Collaborative working in a professional environment')
add_bullet('Self-directed learning and research')

add_page_break()

# CHALLENGES FACED
add_heading_styled('CHALLENGES FACED & SOLUTIONS IMPLEMENTED', level=1)

add_bold_paragraph('Challenge 1: Information Overload')
add_normal_paragraph(
    'The volume of new concepts introduced each day was substantial. Absorbing financial '
    'ratios, chart patterns, and planning frameworks in quick succession was initially '
    'overwhelming.'
)
add_bold_paragraph('Solution:')
add_normal_paragraph(
    'I adopted a practice of summarising key takeaways at the end of each day and '
    'reviewing them before the next session. Creating concise notes with formulas and '
    'frameworks helped in retention and quick reference.'
)

add_bold_paragraph('Challenge 2: Connecting Theory to Practice')
add_normal_paragraph(
    'Initial sessions on portfolio theory and risk-return were abstract. Understanding '
    'how these concepts translate into actual investment decisions required a shift in '
    'thinking.'
)
add_bold_paragraph('Solution:')
add_normal_paragraph(
    'The practical exercises, especially the case study on Day 7, helped bridge this gap. '
    'I also spent extra time exploring real company data on screeners to see how theoretical '
    'ratios looked in practice.'
)

add_bold_paragraph('Challenge 3: Technical Analysis Complexity')
add_normal_paragraph(
    'Interpreting multiple indicators simultaneously and making trading decisions based '
    'on conflicting signals was confusing initially.'
)
add_bold_paragraph('Solution:')
add_normal_paragraph(
    'The trainer\'s advice to focus on confluence — where multiple indicators agree — '
    'simplified the approach. I practised on historical charts to build pattern recognition '
    'before attempting live market analysis.'
)

add_bold_paragraph('Challenge 4: Report Writing Under Time Pressure')
add_normal_paragraph(
    'Compiling a comprehensive internship report while ensuring accuracy, completeness, '
    'and professional formatting within the training period was demanding.'
)
add_bold_paragraph('Solution:')
add_normal_paragraph(
    'I maintained running notes throughout the programme, which significantly reduced '
    'the effort required during final report compilation. Using a structured template '
    'also helped organise content efficiently.'
)

add_page_break()


# PRACTICAL EXPOSURE GAINED
add_heading_styled('PRACTICAL EXPOSURE GAINED', level=1)
add_normal_paragraph(
    'The internship at Fundsweb provided practical exposure across multiple dimensions '
    'of financial services:'
)
add_bullet('Hands-on use of stock screening tools (Screener.in, Tickertape) for filtering investment candidates based on quantitative criteria.')
add_bullet('Real-time chart analysis using TradingView platform with multiple technical indicators and timeframes.')
add_bullet('Construction of goal-based financial plans incorporating inflation adjustment, risk profiling, and product selection.')
add_bullet('Building functional financial calculators that automate routine analytical computations.')
add_bullet('Solving realistic client scenarios involving multiple competing financial goals and constraints.')
add_bullet('Exposure to regulatory frameworks (SEBI, AMFI) governing financial services in India.')
add_bullet('Professional report writing following industry standards for structure, data presentation, and conclusion drawing.')
add_bullet('Understanding of how research desks operate — from data collection through analysis to recommendation generation.')

add_normal_paragraph(
    'This practical exposure has significantly enhanced my confidence in applying '
    'classroom knowledge to real-world financial scenarios. The skills developed during '
    'this programme are directly transferable to roles in equity research, financial '
    'planning, portfolio management, and investment advisory.'
)

add_page_break()

# CONCLUSION
add_heading_styled('CONCLUSION', level=1)
add_normal_paragraph(
    'The 8-day internship training programme at Fundsweb has been a comprehensive and '
    'rewarding experience that has significantly contributed to my professional development '
    'as a Financial Analyst Intern. The programme covered the entire value chain of '
    'investment research — from understanding the industry landscape and company operations '
    'to performing detailed fundamental and technical analysis, and finally communicating '
    'findings through professional reports.'
)
add_normal_paragraph(
    'The structured approach of the training, with each day building upon the previous '
    'one\'s concepts, ensured a smooth learning curve. Starting with organisational '
    'orientation and portfolio basics, progressing through stock analysis and technical '
    'tools, and culminating in practical calculator building and case study solving — '
    'the programme delivered a holistic understanding of the investment research process.'
)
add_normal_paragraph(
    'Key takeaways from this internship include: the importance of goal-based investing, '
    'the complementary nature of fundamental and technical analysis, the critical role of '
    'risk management in both investing and trading, and the necessity of clear '
    'communication in financial services. These learnings will serve as a strong foundation '
    'for my future career in the finance industry.'
)
add_normal_paragraph(
    'I am grateful to Fundsweb for providing this opportunity and to the training team '
    'for their patience, expertise, and commitment to developing the next generation of '
    'financial professionals. This experience has not only enhanced my technical knowledge '
    'but also strengthened my resolve to pursue a career in financial analysis and '
    'investment research.'
)

add_page_break()

# ACKNOWLEDGEMENT
add_heading_styled('ACKNOWLEDGEMENT', level=1)
add_normal_paragraph(
    'I would like to express my sincere gratitude to Fundsweb for providing me with '
    'the opportunity to undertake this internship training programme. The experience '
    'has been invaluable in shaping my understanding of the finance and investment '
    'research industry.'
)
add_normal_paragraph(
    'I am deeply thankful to my reporting manager and the entire training team for their '
    'guidance, mentorship, and willingness to share their industry knowledge and '
    'experience. Their structured approach to training and their patience in addressing '
    'queries made the learning process both effective and enjoyable.'
)
add_normal_paragraph(
    'I also wish to thank my fellow interns for the collaborative learning environment '
    'we created together. The discussions, debates, and shared problem-solving sessions '
    'enriched the overall experience significantly.'
)
add_normal_paragraph(
    'Finally, I am grateful to my college/university for facilitating this internship '
    'opportunity and for providing the academic foundation that enabled me to make the '
    'most of this practical exposure.'
)
add_normal_paragraph(
    'This internship has reinforced my passion for financial analysis and has given me '
    'the confidence and skills to take the next step in my professional journey.'
)

# ========================
# SAVE DOCUMENT
# ========================
output_path = '/projects/sandbox/Hatch/Fundsweb_Internship_Report.docx'
doc.save(output_path)
print(f"Report generated successfully: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
