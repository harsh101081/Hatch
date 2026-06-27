# -*- coding: utf-8 -*-
"""Generate a Financial Services sector comparative business research report (.docx),
mirroring the structure/format of the reference IT Services report."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

NAVY = RGBColor(0x1F, 0x37, 0x64)
GREY = RGBColor(0x44, 0x44, 0x44)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    runs = text.split('\n')
    for i, line in enumerate(runs):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.size = Pt(16)
    return p

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.size = Pt(13)
    return p

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = GREY
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '999999')
    pb.append(bottom); pPr.append(pb)

def make_table(headers, rows, widths=None, header_fill='1F3764'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_text(hdr[i], htext, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=10)
        shade(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def swot_table(strengths, weaknesses, opps, threats):
    t = doc.add_table(rows=2, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ('STRENGTHS', strengths, 'E2EFDA'),
        ('WEAKNESSES', weaknesses, 'FCE4D6'),
        ('OPPORTUNITIES', opps, 'DDEBF7'),
        ('THREATS', threats, 'FBE5E1'),
    ]
    coords = [(0,0),(0,1),(1,0),(1,1)]
    for (title, items, fill), (r,c) in zip(data, coords):
        cell = t.cell(r,c)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = NAVY
        for it in items:
            bp = cell.add_paragraph()
            br = bp.add_run('• ' + it)
            br.font.size = Pt(9)
        shade(cell, fill)
    return t

def page_break():
    doc.add_page_break()

print("scaffolding loaded")


# ============================ TITLE PAGE ============================
def center(text, size, bold=True, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

center("BUSINESS RESEARCH REPORT", 26, color=NAVY, space_before=36, space_after=18)
center("Sector-Based Comparative Company Study", 15, color=GREY, space_after=4)
center("FINANCIAL SERVICES SECTOR", 20, color=NAVY, space_after=28)

companies_title = [
    "1. HDFC Bank Limited",
    "2. ICICI Bank Limited",
    "3. State Bank of India (SBI)",
    "4. Axis Bank Limited",
    "5. Kotak Mahindra Bank Limited",
]
for c in companies_title:
    center(c, 13, bold=True, color=GREY, space_after=4)

doc.add_paragraph().paragraph_format.space_after = Pt(24)

details = [
    ("Student Name:", "Bhavesh Patil"),
    ("Roll Number:", "P2527021"),
    ("College Name:", "Thakur Global Business School"),
    ("PGDM Specialization:", "Finance"),
    ("Date of Submission:", "27/06/2026"),
]
for label, val in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + " "); r1.bold = True; r1.font.size = Pt(12)
    r2 = p.add_run(val); r2.font.size = Pt(12)

page_break()

# ============================ EXECUTIVE SUMMARY ============================
h1("Executive Summary")
body("This research report presents a comprehensive comparative study of five leading institutions in the Indian Financial Services sector: HDFC Bank, ICICI Bank, State Bank of India (SBI), Axis Bank, and Kotak Mahindra Bank. The Financial Services sector is the backbone of the Indian economy, mobilising household savings, allocating capital, enabling payments, and funding the credit needs of individuals, businesses, and government. Banking alone accounts for the largest share of India's financial system, and the five institutions studied here collectively command a dominant share of system deposits, advances, and market capitalisation.")
body("India's banking sector has entered one of its healthiest phases in over a decade. System-wide gross non-performing assets fell to multi-year lows, credit growth remained robust in double digits, and the listed private and public-sector banks delivered record profitability in FY2024-25. Together, these five institutions represent the core of India's formal banking system, blending the scale of public-sector banking with the efficiency and digital innovation of leading private banks.")
body("The major findings of this report include:")
for b in [
    "SBI is the undisputed market leader by balance-sheet size, deposits, branch network, and customer base, and reported a record standalone net profit of approximately Rs 70,901 crore in FY25 (up ~16% YoY) with asset quality at a multi-year best (GNPA ~1.82%).",
    "HDFC Bank is India's largest private-sector bank and, following its 2023 merger with parent HDFC Ltd, one of the largest banks in the world by market capitalisation. It posted standalone net profit of approximately Rs 67,347 crore in FY25 with industry-leading asset quality (GNPA ~1.33%).",
    "ICICI Bank delivered sector-leading profitability and returns, with standalone net profit of approximately Rs 47,227 crore (up ~15.5% YoY), one of the highest return-on-assets among large banks, and a strong digital franchise.",
    "Axis Bank, the third-largest private bank, reported net profit of approximately Rs 26,373 crore with healthy margins (FY25 NIM ~3.98%) and continued improvement in its deposit franchise and asset quality.",
    "Kotak Mahindra Bank, the most conservatively run of the five, delivered standalone net profit of approximately Rs 16,450 crore, the highest net interest margin (~5%) among the group, and a diversified financial-services group spanning broking, asset management, and insurance.",
]:
    bullet(b)
body("All five institutions are investing aggressively in digital banking, payments, data analytics, and artificial intelligence as key growth and efficiency vectors. Deposit mobilisation in a high-interest-rate environment, net interest margin protection, asset-quality discipline in unsecured retail lending, and regulatory compliance remain common sectoral challenges.")
hr()
page_break()
print("title + exec summary loaded")


# ============================ 1. INTRODUCTION ============================
h1("1. Introduction to the Financial Services Sector")
h2("1.1 Industry Overview")
h3("What is the Financial Services Sector?")
body("The Financial Services sector encompasses the broad range of institutions and activities involved in the management of money. It includes commercial and retail banking, lending and credit, deposit mobilisation, payments and settlements, wealth and asset management, insurance, broking, and investment banking. In the Indian context, banks form the dominant pillar of this sector, intermediating between savers and borrowers, powering the payments ecosystem, and serving as the primary channel for the transmission of monetary policy. The five companies studied in this report are all scheduled commercial banks, anchoring large financial-services groups that extend into insurance, mutual funds, broking, and non-banking finance.")
h3("Economic Importance")
body("The financial services sector is central to India's economic growth. The banking system channels household savings into productive investment, funds infrastructure and industry, and extends credit to micro, small, and medium enterprises as well as to retail consumers. A well-capitalised, well-governed banking system is a prerequisite for sustained GDP growth. India's formal banking system has expanded financial inclusion dramatically over the past decade, with hundreds of millions of new accounts opened under the Jan Dhan programme and the Unified Payments Interface (UPI) processing billions of transactions every month, making India one of the largest real-time digital-payments markets in the world.")
h3("Current Market Size and Growth Trends")
body("Indian banks have been growing their balance sheets at a healthy double-digit pace, with system credit growth in the low-to-mid teens and deposit growth gradually catching up. System-wide asset quality improved markedly, with gross non-performing assets of scheduled commercial banks falling to multi-year lows, supported by strong recoveries, prudent provisioning, and disciplined underwriting. Profitability across the listed banking universe reached record highs in FY2024-25, with leading private banks reporting return-on-assets near or above 1.8% and the largest public-sector bank crossing the 1.1% mark. The sector's medium-term growth is underpinned by India's favourable demographics, rising per-capita income, formalisation of the economy, and the rapid digitisation of financial services.")
h3("Future Opportunities")
body("The sector presents substantial growth opportunities driven by: (a) deepening retail credit penetration in mortgages, vehicle loans, and consumer finance; (b) the formalisation and digitisation of MSME lending through account aggregators and cash-flow-based underwriting; (c) the expansion of digital payments, co-lending, and embedded finance; (d) the growth of wealth management and insurance as household financial savings shift toward market-linked products; and (e) financial inclusion in semi-urban and rural India through phygital distribution and business correspondents.")
h3("Major Challenges")
body("The sector faces a number of structural and cyclical challenges including: intense competition for low-cost deposits (CASA) in a high-interest-rate environment, which pressures net interest margins; rising stress in unsecured retail and microfinance lending; the need for continuous investment in cybersecurity and technology resilience; evolving regulatory expectations on capital, liquidity, and governance; and the disruptive entry of fintech players and Big Tech into payments and lending.")
h3("Role of Technology")
body("Technology is reshaping every aspect of banking. The five institutions in this study have all pivoted aggressively toward digital-first banking through mobile apps, internet banking, UPI, and API-led open banking. Data analytics and artificial intelligence are being deployed for credit underwriting, fraud detection, personalisation, and customer service automation through chatbots and virtual assistants. Core-banking modernisation, cloud adoption, and robust cybersecurity have become board-level priorities. The adoption of AI is moving banks from product-centric to customer-centric models, enabling hyper-personalised offers and real-time risk management.")
h3("Government Regulations")
body("Key regulatory frameworks affecting the sector include: (i) the Banking Regulation Act, 1949, and the Reserve Bank of India Act, 1934, governing licensing, capital, and prudential norms; (ii) RBI guidelines on capital adequacy (Basel III), priority-sector lending, asset classification, and provisioning; (iii) the Insolvency and Bankruptcy Code, 2016, which transformed corporate-loan recoveries; (iv) the Digital Personal Data Protection Act, 2023, governing customer-data privacy; (v) RBI's digital-lending guidelines and rules on co-lending and outsourcing; and (vi) SEBI and IRDAI regulations applicable to the banks' broking, asset-management, and insurance subsidiaries.")

h2("1.2 Sector Analysis")
h3("Major Players")
body("The Indian banking sector is dominated by a mix of large public-sector and private-sector banks. State Bank of India is the largest bank by every scale metric. Among private banks, HDFC Bank, ICICI Bank, Axis Bank, and Kotak Mahindra Bank form the leading tier, collectively referred to as India's leading private banks. The next tier includes IndusInd Bank, IDBI Bank, Yes Bank, and the larger public-sector banks such as Bank of Baroda, Punjab National Bank, and Canara Bank. Small finance banks, payments banks, NBFCs, and fintech players compete in specific niches such as microfinance, consumer lending, and digital payments.")
h3("Competition Factors")
body("Competition is driven by: scale and reach of the deposit franchise (particularly low-cost CASA deposits); cost of funds and net interest margin; breadth and quality of the loan book; digital capability and customer experience; asset-quality discipline and risk management; fee-income diversification (cards, wealth, transaction banking); brand trust; and the strength of the broader financial-services group across insurance, asset management, and broking.")
h3("Sector Status")
body("The Indian banking sector is in a strong growth and consolidation phase. Balance sheets are well-capitalised, asset quality is at a multi-year best, and profitability is at record levels. While net interest margins face near-term pressure from a potential rate-cut cycle and competition for deposits, the sector's long-term fundamentals — driven by credit penetration, digitisation, and financial inclusion — remain firmly positive.")
hr()
page_break()

# ============================ 2. COMPANY SELECTION ============================
h1("2. Company Selection")
body("The five companies selected for this comparative study all operate within the Indian Financial Services sector as scheduled commercial banks listed on the BSE and NSE. Together, they represent the largest banks in India by deposits, advances, and market capitalisation, spanning both the public-sector leader and the four leading private-sector banks, as of FY2024-25.")
make_table(
    ["Sr.", "Company", "HQ", "Founded", "Founder / Origin", "Net Profit (FY25)", "Employees (approx.)"],
    [
        ["1", "State Bank of India", "Mumbai", "1955*", "Government of India (origins 1806)", "~Rs 70,901 cr", "~2,36,000"],
        ["2", "HDFC Bank", "Mumbai", "1994", "Housing Development Finance Corp.", "~Rs 67,347 cr", "~2,14,000"],
        ["3", "ICICI Bank", "Mumbai", "1994", "ICICI Limited", "~Rs 47,227 cr", "~1,33,000"],
        ["4", "Axis Bank", "Mumbai", "1993", "UTI & insurance institutions", "~Rs 26,373 cr", "~1,05,000"],
        ["5", "Kotak Mahindra Bank", "Mumbai", "1985**", "Uday Kotak", "~Rs 16,450 cr", "~74,000"],
    ],
    widths=[0.4, 1.6, 0.8, 0.8, 1.6, 1.1, 1.0],
)
body("*State Bank of India was constituted in its present form in 1955 under the State Bank of India Act, but its origins trace to the Bank of Calcutta founded in 1806 — making it the oldest commercial bank in the Indian subcontinent. **Kotak Mahindra was founded in 1985 as Kotak Mahindra Finance Limited and became the first NBFC in India to be converted into a commercial bank, receiving its banking licence in 2003. Net-profit figures are standalone for FY2024-25.")
hr()
page_break()
print("sections 1-2 loaded")


# ============================ 3. COMPANY PROFILE ANALYSIS ============================
h1("3. Company Profile Analysis")

# ---------------- 3.1 HDFC BANK ----------------
h2("3.1 HDFC Bank")
h3("Basic Company Understanding")
body("HDFC Bank is India's largest private-sector bank, headquartered in Mumbai. It was incorporated in 1994 as part of the liberalisation of the Indian banking sector, promoted by Housing Development Finance Corporation (HDFC Ltd), India's pioneering housing-finance institution. In July 2023, HDFC Bank completed a landmark merger with its parent HDFC Ltd — the largest transaction in Indian corporate history — creating a financial powerhouse with a combined balance sheet that ranks among the largest banks globally by market capitalisation.")
body("HDFC Bank offers a full suite of banking and financial services spanning retail banking, wholesale (corporate) banking, treasury operations, and a wide range of digital products. Its subsidiaries include HDB Financial Services (NBFC), HDFC Securities (broking), and HDFC Life and HDFC AMC within the broader group. Target customers span retail consumers, MSMEs, large corporates, and the government.")
body("Vision: 'To be a world-class Indian bank.' The bank pursues two core objectives — to be the preferred provider of banking services for target retail and wholesale customer segments, and to achieve healthy growth in profitability consistent with the bank's risk appetite. Core Values: Operational Excellence, Customer Focus, Product Leadership, People, and Sustainability. The bank is led by MD & CEO Sasidhar Jagdishan and Chairman Atanu Chakraborty.")
h3("Business Model")
body("HDFC Bank operates a universal banking model serving retail and wholesale customers. Its business model rests on a large, low-cost deposit franchise funding a diversified, high-quality loan book across mortgages, retail loans, MSME lending, and corporate credit. Net interest income is the primary revenue driver, supplemented by a strong fee-income engine from cards, payments, third-party product distribution, and transaction banking.")
body("Key resources include a workforce of approximately 2.14 lakh employees, one of the largest branch and ATM networks in the private sector spanning urban, semi-urban, and rural India, a market-leading credit-card and payments franchise, and a robust digital banking platform. Following the merger, mortgages have become a much larger part of the loan mix, deepening the bank's secured-lending profile.")
h3("Revenue Generation")
body("In FY2024-25, HDFC Bank reported total income of approximately Rs 4,70,916 crore and standalone net profit of approximately Rs 67,347 crore, with consolidated net profit of approximately Rs 70,792 crore. Gross advances stood at approximately Rs 26.4 lakh crore as of March 31, 2025, and average deposits grew nearly 16% year-on-year, with the bank deliberately allowing deposit growth to outpace loan growth to bring down its post-merger credit-deposit ratio.")
body("Revenue is generated through net interest income on its loan book, fee and commission income, and treasury operations. Net interest margin for the year was approximately 3.5% on total assets. The Board recommended a dividend of Rs 22 per share for FY25, reflecting strong capital generation.")
h3("Financial Analysis")
body("HDFC Bank maintains best-in-class financial metrics. Asset quality is the strongest among large banks, with gross NPA of approximately 1.33% and net NPA of approximately 0.43% as of March 31, 2025. The capital adequacy ratio stood at a healthy 19.6%, with CET1 at 17.2%. Return on assets remained around 1.8% and return on equity in the mid-teens. The bank is well-provisioned and carries a strong floating-provision buffer, underscoring conservative risk management.")
body("Key Financial Ratios (FY25 estimates): Net Interest Margin ~3.5%; Gross NPA ~1.33%; Net NPA ~0.43%; Capital Adequacy ~19.6%; ROA ~1.8%; ROE ~14%; Cost-to-Income ~40%.")
h3("Marketing Strategy")
body("HDFC Bank invests heavily in brand building, consistently ranking among India's most valuable brands. Its marketing spans mass-media campaigns, financial-literacy initiatives, digital and social-media engagement, and large-scale festive credit-card and EMI promotions. The bank leverages co-branded card partnerships with airlines, retailers, and e-commerce platforms to acquire and retain affluent customers. Its USP is the combination of trust, scale, and a market-leading payments and cards franchise delivered with consistent service quality.")
h3("Sales Strategy")
body("HDFC Bank uses a multi-channel distribution model combining an extensive branch network, direct sales agents, relationship managers for affluent and corporate clients, and a powerful digital acquisition engine. Cross-selling across the product suite — savings accounts, cards, loans, insurance, and investments — is central to its strategy, supported by deep customer analytics and a relationship-based banking approach.")
h3("HR & Organisational Culture")
body("HDFC Bank employs approximately 2.14 lakh people and is one of India's largest private-sector employers. It recruits extensively from campuses and the open market, with structured training academies for sales, credit, and leadership development. The culture emphasises performance, process discipline, customer focus, and integrity. The bank has invested in employee well-being, diversity, and digital upskilling as it transforms into a technology-led organisation.")
h3("Technology & Innovation")
body("HDFC Bank has embarked on a comprehensive 'Digital Factory' and technology-transformation programme, modernising its core systems, strengthening cloud infrastructure, and building new digital products. It has invested heavily in cybersecurity and resilience following earlier regulatory scrutiny of its technology outages. AI and machine learning are deployed for fraud detection, credit decisioning, and personalised customer engagement, with conversational banking and a revamped mobile app at the centre of its digital strategy.")
h3("Customer Analysis")
body("HDFC Bank serves a vast customer base of retail consumers, small businesses, and large corporates. It is the market leader in credit cards and a leading player in payments, auto loans, and personal loans. Customer experience is managed through dedicated relationship teams for premium segments, a wide branch network, and 24x7 digital and phone banking. The bank's strong brand trust and service consistency drive high customer retention and cross-sell rates.")
h3("Competitor Analysis")
body("HDFC Bank's primary competitors are ICICI Bank and Axis Bank among private banks, and SBI in the broader market. Its competitive advantages include the largest private-sector balance sheet, best-in-class asset quality, a dominant cards and payments franchise, and the deep mortgage capabilities gained through the HDFC Ltd merger. Its key challenge is to drive deposit mobilisation at scale to fund growth while protecting margins.")

# ---------------- 3.2 ICICI BANK ----------------
h2("3.2 ICICI Bank")
h3("Basic Company Understanding")
body("ICICI Bank is India's second-largest private-sector bank and among the most profitable large banks in the country, headquartered in Mumbai. It was incorporated in 1994 as a subsidiary of ICICI Limited, a development-finance institution, and in 2002 ICICI Limited reverse-merged into the bank. ICICI Bank is listed on the NSE and BSE in India and on the NYSE (IBN) in the United States, and serves customers across India and key international markets.")
body("ICICI Bank offers a comprehensive range of banking products and financial services across retail, SME, and corporate banking, treasury, and digital channels. Its subsidiaries include ICICI Prudential Life Insurance, ICICI Lombard General Insurance, ICICI Prudential AMC, and ICICI Securities, making it one of India's most diversified financial-services groups. The bank is led by MD & CEO Sandeep Bakhshi.")
body("Vision: 'To be the trusted financial-services provider of choice for our customers, delivering sustainable returns to shareholders.' The bank's strategic philosophy centres on the principles of 'Fair to Customer, Fair to Bank' and maximising risk-calibrated core operating profit. Core Values include customer-centricity, integrity, and a strong risk culture.")
h3("Business Model")
body("ICICI Bank follows a universal banking model with a sharp focus on risk-calibrated profitable growth. Its '360-degree customer-centric' approach organises the bank around customer ecosystems rather than products, leveraging data and digital platforms to deepen relationships. Revenue is driven by net interest income, a granular fee-income stream, and the strong performance of its insurance and securities subsidiaries.")
body("Key resources include a workforce of approximately 1.33 lakh employees, an extensive branch and ATM network, market-leading digital platforms (iMobile Pay, InstaBIZ, and the API banking suite), and a powerful group ecosystem spanning life and general insurance, asset management, and broking. The bank's Business Banking segment, serving SMEs, has been a standout growth engine.")
h3("Revenue Generation")
body("In FY2024-25, ICICI Bank reported standalone net profit of approximately Rs 47,227 crore, up about 15.5% year-on-year, making it one of the most profitable private banks in India. Profit before tax excluding treasury — the bank's preferred measure of core profitability — grew at a healthy double-digit pace. Revenue is generated through net interest income on a well-diversified loan book skewed toward retail and business banking, complemented by strong fee income from cards, payments, and transaction banking.")
body("The bank maintains one of the highest net interest margins among large banks at approximately 4.4%, supported by a strong low-cost deposit franchise and a favourable loan mix. Its consolidated profit, including subsidiaries, is materially higher than the standalone figure.")
h3("Financial Analysis")
body("ICICI Bank delivers sector-leading returns. Return on assets is among the highest in the large-bank universe at approximately 2.4%, and return on equity is around 18%. Asset quality is strong, with gross NPA of approximately 1.67% and net NPA of approximately 0.39%, supported by a high provision-coverage ratio. The bank is well-capitalised with a capital adequacy ratio comfortably above regulatory requirements.")
body("Key Financial Ratios (FY25 estimates): Net Interest Margin ~4.4%; Gross NPA ~1.67%; Net NPA ~0.39%; ROA ~2.4%; ROE ~18%; Provision Coverage ~78%; Cost-to-Income ~38%.")
h3("Marketing Strategy")
body("ICICI Bank positions itself as a technology-led, customer-centric bank. Its marketing emphasises digital convenience, with campaigns built around its award-winning mobile and internet-banking platforms. The bank leverages co-branded cards, sponsorships, and ecosystem partnerships to acquire customers, and uses data-driven personalisation to deepen relationships. Its USP is the combination of high profitability, robust risk management, and digital leadership.")
h3("HR & Organisational Culture")
body("ICICI Bank employs approximately 1.33 lakh people and rationalised its headcount in FY25 as productivity improved through automation. The culture is performance-driven, meritocratic, and increasingly digital-first, with a strong emphasis on risk discipline embedded across the organisation. The bank invests in structured training, leadership development, and a 'one-bank, one-team' collaborative ethos that breaks down product silos.")
h3("Technology & Innovation")
body("ICICI Bank is widely regarded as a digital-banking leader. Its iMobile Pay app offers a full suite of banking and payment services to customers of any bank, while InstaBIZ serves business customers and its API banking platform powers embedded-finance partnerships. The bank applies AI and analytics across underwriting, fraud management, collections, and hyper-personalised customer engagement, and was an early mover in cloud adoption and decisioning automation.")
h3("Customer Analysis")
body("ICICI Bank serves retail consumers, small and medium enterprises, and large corporates through a 360-degree ecosystem approach. Its rapidly growing Business Banking portfolio, retail mortgages, and credit cards are key franchises. Customer experience is delivered through a blend of digital self-service, relationship managers, and branches, with strong digital adoption metrics and high customer-satisfaction scores.")

# ---------------- 3.3 SBI ----------------
h2("3.3 State Bank of India (SBI)")
h3("Basic Company Understanding")
body("State Bank of India is the largest commercial bank in India by every scale metric — assets, deposits, advances, branches, and customers — and is majority-owned by the Government of India. Headquartered in Mumbai, SBI traces its origins to the Bank of Calcutta founded in 1806, making it the oldest commercial bank in the Indian subcontinent. It was constituted in its present form in 1955 under the State Bank of India Act and has since served as the backbone of India's banking system and a key instrument of financial inclusion.")
body("SBI offers the full spectrum of banking and financial services to individuals, businesses, corporates, and the government across India and through a wide international network. Its subsidiaries and group companies include SBI Life Insurance, SBI Cards, SBI Mutual Fund (SBI Funds Management), SBI General Insurance, and SBI Capital Markets — collectively one of India's largest financial-services ecosystems. The bank is led by Chairman C. S. Setty.")
body("Mission: 'To be the bank of choice for a transforming India.' Vision: 'Be the most preferred and most valued bank with global standards.' Core Values, encapsulated in the acronym S.T.E.P.S., are Service, Transparency, Ethics, Politeness, and Sustainability.")
h3("Business Model")
body("SBI operates a universal banking model at unmatched national scale. Its business model is built on an enormous, granular, low-cost deposit base — among the highest CASA ratios in the system — funding a diversified loan book spanning retail, agriculture, MSME, and corporate lending. The bank plays a central role in government banking, social-sector schemes, and infrastructure financing, while its subsidiaries contribute significant value through insurance, cards, and asset management.")
body("Key resources include a workforce of approximately 2.36 lakh employees, the largest branch and ATM network in the country reaching deep into rural India, and the market-leading YONO digital platform. SBI's scale, sovereign backing, and trust make it the default banking choice for a vast cross-section of the Indian population.")
h3("Revenue Generation")
body("In FY2024-25, SBI reported a record standalone net profit of approximately Rs 70,901 crore, up about 16% year-on-year — the highest ever for any Indian bank. Revenue is generated through net interest income on its massive loan book, fee income from transaction banking and government business, and treasury operations, supplemented by strong dividend and value contributions from its insurance, cards, and asset-management subsidiaries.")
body("The bank's total business surpassed historic milestones, with deposits of approximately Rs 53.8 lakh crore. Its scale advantage in low-cost deposits gives it a structural funding edge, even though its net interest margin is lower than the leading private banks.")
h3("Financial Analysis")
body("SBI's FY25 results marked a structural improvement in profitability and asset quality. Return on assets improved to approximately 1.10% and return on equity rose to around 19-20%. Asset quality reached a multi-year best, with gross NPA of approximately 1.82% and net NPA of approximately 0.47%, supported by strong recoveries and prudent provisioning. The bank remains adequately capitalised and continues to raise capital as needed to fund growth.")
body("Key Financial Ratios (FY25 estimates): Net Interest Margin ~3.0%; Gross NPA ~1.82%; Net NPA ~0.47%; ROA ~1.10%; ROE ~19-20%; CASA Ratio ~40%; Provision Coverage ~74%.")
h3("Marketing Strategy")
body("As a household name with over two centuries of heritage, SBI enjoys unparalleled brand trust. Its marketing focuses on accessibility, financial inclusion, and the reliability of the national bank, with campaigns spanning mass media, regional languages, and rural outreach. SBI leverages its YONO platform and SBI Card brand to engage younger, digitally native customers, balancing its legacy image with a modern digital identity.")
h3("HR & Organisational Culture")
body("SBI employs approximately 2.36 lakh people, making it one of the largest employers in India. It recruits through highly competitive nationwide examinations and invests in extensive training through its apex training institutes. The culture combines public-sector values of service and stability with a growing emphasis on performance, digital skills, and customer experience. Succession is structured and merit-based, with leaders typically rising through decades of service across diverse geographies.")
h3("Technology & Innovation")
body("SBI's YONO ('You Only Need One') platform is one of the world's largest digital-banking ecosystems, offering banking, investments, insurance, and shopping in a single app, with registrations in the tens of crores. The bank has digitised account opening, lending, and servicing at scale, and is investing in AI, analytics, and cybersecurity to manage its enormous transaction volumes. SBI processes a very high share of transactions through digital and alternate channels, reflecting deep digital adoption across its customer base.")
h3("Customer Analysis")
body("SBI serves the broadest customer base of any Indian bank, from rural Jan Dhan account holders to large corporates and the government. Its anchor strengths are deep penetration in semi-urban and rural India, dominance in government and salary banking, and trust built over generations. Customer service is delivered through its vast branch network, business correspondents, the YONO app, and a large contact-centre operation.")

# ---------------- 3.4 AXIS BANK ----------------
h2("3.4 Axis Bank")
h3("Basic Company Understanding")
body("Axis Bank is India's third-largest private-sector bank, headquartered in Mumbai. It was founded in 1993 as UTI Bank — one of the first new private banks established after liberalisation — promoted by the Unit Trust of India along with leading insurance institutions including LIC and GIC. It was renamed Axis Bank in 2007. The bank serves retail, SME, and corporate customers across a wide national network and select international locations.")
body("Axis Bank offers a comprehensive suite of banking products spanning retail banking, corporate and commercial banking, treasury, and digital services. Its subsidiaries include Axis Securities, Axis Asset Management, Axis Finance, and Axis Capital, and it holds a significant stake in Max Life Insurance through a strategic partnership. The bank is led by MD & CEO Amitabh Chaudhry.")
body("Vision: 'To be the preferred financial-services provider, excelling in customer service delivery through insight, empowered employees, and smart use of technology.' Core Values: Customer Centricity, Ethics, Transparency, Teamwork, and Ownership.")
h3("Business Model")
body("Axis Bank operates a universal banking model with a strategic focus on the 'GPS' framework — Growth, Profitability, and Sustainability. Its business model balances a granular retail franchise with strong corporate and transaction-banking businesses. The bank has steadily improved its deposit quality, fee-income mix, and digital capabilities, while maintaining a disciplined approach to risk and capital.")
body("Key resources include a workforce of approximately 1.05 lakh employees, a wide branch and ATM network, a strong cards and payments franchise, and the market-leading Axis Mobile and 'open by Axis Bank' digital platforms. The acquisition of Citibank's India consumer-banking business strengthened its affluent-customer and cards franchise.")
h3("Revenue Generation")
body("In FY2024-25, Axis Bank reported net profit of approximately Rs 26,373 crore, up about 6% year-on-year. Revenue is driven by net interest income, with a full-year net interest margin of approximately 3.98%, supplemented by a granular fee-income stream in which retail fees constitute the bulk of the total. Core operating profit grew at a healthy double-digit pace, reflecting improving operating leverage.")
body("The bank's fee income grew strongly, with granular (non-treasury) fees forming the vast majority of total fees, underscoring the quality and stability of its earnings. Consolidated return on assets and return on equity reached approximately 1.77% and 16.89% respectively.")
h3("Financial Analysis")
body("Axis Bank's financial profile has strengthened considerably. Asset quality improved to a gross NPA of approximately 1.28% and net NPA of approximately 0.33%, with a provision-coverage ratio around 75%. Cost-to-assets declined as the bank drove operating efficiency. The bank is well-capitalised and has focused on building a more stable, lower-cost deposit base to support sustainable growth.")
body("Key Financial Ratios (FY25 estimates): Net Interest Margin ~3.98%; Gross NPA ~1.28%; Net NPA ~0.33%; ROA ~1.77%; ROE ~16.89%; Provision Coverage ~75%; Cost-to-Assets ~2.46%.")
h3("Marketing Strategy")
body("Axis Bank markets itself as a modern, digitally progressive private bank with the tagline philosophy of helping customers progress ('Dil Se Open'). Its campaigns emphasise digital convenience, premium credit cards, and lifestyle banking for the affluent and aspirational middle class. The Citibank consumer-business acquisition added a premium customer base and marquee co-branded card portfolios, strengthening its position in the high-value cards segment.")
h3("HR & Organisational Culture")
body("Axis Bank employs approximately 1.05 lakh people. The culture under its current leadership emphasises customer obsession, ethics, and execution discipline through the GPS framework. The bank invests in leadership development, digital and analytics talent, and diversity and inclusion. Integration of the acquired Citi consumer-banking workforce has been a key recent organisational priority.")
h3("Technology & Innovation")
body("Axis Bank has invested significantly in digital transformation through its in-house digital banking unit. Its 'open by Axis Bank' and Axis Mobile apps are consistently ranked among the best in the industry. The bank applies AI and analytics across customer acquisition, underwriting, cross-sell, and fraud management, and has built API-led partnerships to power embedded finance and co-lending. Cloud adoption and cybersecurity are central to its technology roadmap.")
h3("Customer Analysis")
body("Axis Bank serves retail consumers, small businesses, and large corporates, with particular strength in cards, retail lending, and transaction banking. The Citibank acquisition deepened its affluent and premium-customer franchise. Customer experience is delivered through a strong digital platform, branches, and relationship managers, with rising digital-transaction adoption and improving customer-satisfaction metrics.")

# ---------------- 3.5 KOTAK MAHINDRA BANK ----------------
h2("3.5 Kotak Mahindra Bank")
h3("Basic Company Understanding")
body("Kotak Mahindra Bank is one of India's leading private-sector banks and the flagship of the Kotak Mahindra group, headquartered in Mumbai. It was founded in 1985 as Kotak Mahindra Finance Limited by Uday Kotak and became, in 2003, the first non-banking finance company in India to be converted into a full-fledged commercial bank. Kotak is renowned for its conservative risk culture, capital discipline, and consistently high profitability.")
body("Kotak Mahindra Bank offers a full range of banking and financial services, and anchors one of India's most diversified financial-services groups spanning broking (Kotak Securities), asset management (Kotak Mahindra AMC), life and general insurance, and investment banking (Kotak Mahindra Capital). The bank is led by MD & CEO Ashok Vaswani, with founder Uday Kotak continuing as a significant shareholder and non-executive director.")
body("Vision: 'To be the most trusted financial-services partner, helping customers achieve their financial goals.' The group's enduring philosophy emphasises prudence, customer trust, and 'concentrated India, deliver global' — building deep capabilities in the Indian market. Core Values centre on integrity, customer-first thinking, and disciplined risk management.")
h3("Business Model")
body("Kotak Mahindra Bank operates a universal banking model distinguished by a conservative, profitability-first approach. Rather than chasing scale, Kotak prioritises risk-adjusted returns, maintaining one of the highest net interest margins and lowest cost-of-credit profiles in the industry. The integrated group model allows the bank to capture value across the customer's entire financial life — banking, investing, protecting, and borrowing.")
body("Key resources include a workforce of approximately 74,000 employees (with the wider group employing more), a focused branch network concentrated in high-value urban and semi-urban markets, the acclaimed Kotak 811 digital-banking platform, and a powerful group ecosystem across capital markets, asset management, and insurance. Group assets under management have grown strongly across equity and debt segments.")
h3("Revenue Generation")
body("In FY2024-25, Kotak Mahindra Bank reported standalone total income of approximately Rs 1,06,902 crore and standalone net profit of approximately Rs 16,450 crore, up from approximately Rs 13,782 crore in FY24. Consolidated net profit, including its subsidiaries, was materially higher at approximately Rs 22,126 crore, reflecting the strong contribution of the broking, asset-management, and insurance businesses.")
body("Revenue is driven by net interest income — supported by the highest net interest margin among the five banks at approximately 5% — and a healthy CASA ratio of around 43%. Fee income, capital-markets income, and subsidiary profits provide significant earnings diversification. An exceptional gain from the divestment of a majority stake in Kotak General Insurance boosted FY25 profitability.")
h3("Financial Analysis")
body("Kotak Mahindra Bank's financial profile is among the strongest in the sector. Asset quality is excellent, with gross NPA of approximately 1.42% and net NPA of approximately 0.31%, and a provision-coverage ratio around 78%. The bank is exceptionally well-capitalised, with one of the highest capital-adequacy ratios in the industry, giving it substantial headroom for growth. Return on assets is consistently around 2%, reflecting its profitability-first model.")
body("Key Financial Ratios (FY25 estimates): Net Interest Margin ~5%; Gross NPA ~1.42%; Net NPA ~0.31%; CASA Ratio ~43%; ROA ~2%; ROE ~13-14%; Provision Coverage ~78%.")
h3("Marketing Strategy")
body("Kotak markets itself on trust, prudence, and a premium, customer-first experience. Its flagship Kotak 811 digital-banking proposition — a zero-balance, fully digital account — has been a powerful customer-acquisition engine for younger and digitally native customers. The bank emphasises its integrated group offering, allowing customers to bank, invest, and protect within a single trusted ecosystem. Its brand is associated with stability and disciplined wealth creation.")
h3("HR & Organisational Culture")
body("Kotak Mahindra Bank employs approximately 74,000 people. The culture reflects the founder's emphasis on prudence, integrity, and long-term thinking, balanced with the need to build a more agile, digital, and growth-oriented organisation under new leadership. The bank invests in talent across digital, technology, and risk functions, and the leadership transition from founder Uday Kotak to CEO Ashok Vaswani has been a defining organisational milestone.")
h3("Technology & Innovation")
body("Kotak has invested significantly in digital banking, led by the Kotak 811 platform and a feature-rich mobile-banking app. The bank applies analytics and AI across onboarding, lending, and customer engagement, and has worked to strengthen its technology resilience and cybersecurity following earlier regulatory action that temporarily restricted new digital onboarding and credit-card issuance. Remediation and core-systems modernisation have been key technology priorities, with the restrictions subsequently addressed.")
h3("Customer Analysis")
body("Kotak Mahindra Bank serves affluent and mass-affluent retail customers, businesses, and corporates, with particular strength in wealth management, capital markets, and premium banking. Its integrated group model allows it to serve the full financial-services needs of high-value customers. Customer experience is delivered through digital platforms, a focused branch network, and relationship managers, with strong cross-sell across the group's products.")
hr()
page_break()
print("section 3 loaded")


# ============================ 4. SWOT ANALYSIS ============================
h1("4. SWOT Analysis - All Five Companies")

h2("4.1 HDFC Bank - SWOT Analysis")
swot_table(
    ["Largest private-sector bank by assets & market cap",
     "Best-in-class asset quality (GNPA ~1.33%)",
     "Dominant credit-card & payments franchise",
     "Strong capital adequacy (~19.6%)",
     "Deep mortgage capability post HDFC Ltd merger",
     "Trusted brand & consistent execution"],
    ["High credit-deposit ratio post-merger",
     "Margin pressure while rebuilding deposit base",
     "Integration complexity of the mega-merger",
     "Past regulatory scrutiny on tech outages",
     "Slower near-term loan growth by design"],
    ["Cross-sell of mortgages to deposit customers",
     "Rural & semi-urban deposit mobilisation",
     "Wealth management for affluent segment",
     "Digital lending & embedded finance",
     "Co-lending partnerships with NBFCs"],
    ["Intense competition for low-cost deposits",
     "Net interest margin compression in rate-cut cycle",
     "Rising stress in unsecured retail credit",
     "Fintech disruption in payments",
     "Cybersecurity & technology-resilience risk"],
)

h2("4.2 ICICI Bank - SWOT Analysis")
swot_table(
    ["Sector-leading profitability (ROA ~2.4%)",
     "High net interest margin (~4.4%)",
     "Digital-banking leadership (iMobile Pay, API banking)",
     "Strong, diversified group (insurance, AMC, broking)",
     "Robust risk culture & strong provisioning",
     "Fast-growing Business Banking franchise"],
    ["Past legacy of corporate-loan stress (now resolved)",
     "High dependence on retail & business banking growth",
     "Exposure to unsecured retail cycle",
     "Headcount rationalisation execution risk"],
    ["Deepening SME & business-banking penetration",
     "Ecosystem & embedded-finance partnerships",
     "Wealth & affluent-banking expansion",
     "Cross-sell across group subsidiaries"],
    ["Competition for deposits & talent",
     "Margin pressure in a falling-rate environment",
     "Asset-quality risk in unsecured lending",
     "Regulatory changes on fees & lending",
     "Cyber & data-security threats"],
)

h2("4.3 State Bank of India - SWOT Analysis")
swot_table(
    ["Largest bank by assets, deposits & branches",
     "Record net profit (~Rs 70,901 cr in FY25)",
     "Unmatched low-cost deposit & CASA franchise",
     "Sovereign backing & deep public trust",
     "YONO - one of the world's largest digital platforms",
     "Strong subsidiaries (SBI Life, SBI Cards, SBI MF)"],
    ["Lower NIM & ROA than leading private banks",
     "Large workforce & legacy cost structure",
     "Government-banking & social-obligation burden",
     "Periodic need for capital raising",
     "Slower decision-making due to scale & ownership"],
    ["Cross-sell to its vast customer base",
     "Monetisation of subsidiary stakes",
     "Rural & MSME credit expansion",
     "Digital-led cost efficiency via YONO",
     "Infrastructure & corporate-credit revival"],
    ["Margin pressure in a rate-cut cycle",
     "Asset-quality risk in agriculture & SME books",
     "Competition from agile private banks & fintechs",
     "Government-policy & priority-sector obligations",
     "Cybersecurity at massive transaction scale"],
)

h2("4.4 Axis Bank - SWOT Analysis")
swot_table(
    ["Third-largest private bank with national reach",
     "Healthy margins (FY25 NIM ~3.98%)",
     "Improved asset quality (GNPA ~1.28%)",
     "Strong cards & retail-fee franchise",
     "Citibank consumer-business acquisition",
     "Best-in-class digital apps ('open by Axis')"],
    ["Lower CASA ratio than top private peers",
     "Historically more volatile earnings",
     "Integration of acquired Citi portfolio",
     "Lower scale than HDFC & ICICI"],
    ["Affluent & premium-cards growth via Citi base",
     "Granular deposit & fee-income expansion",
     "Digital lending & partnership banking",
     "Wealth & transaction-banking cross-sell"],
    ["Competition for deposits & affluent customers",
     "Margin pressure & rising cost of funds",
     "Unsecured-lending asset-quality cycle",
     "Fintech disruption in payments & lending",
     "Cyber & operational-risk exposure"],
)

h2("4.5 Kotak Mahindra Bank - SWOT Analysis")
swot_table(
    ["Highest NIM among the five (~5%)",
     "Conservative risk culture & strong capital",
     "Excellent asset quality (GNPA ~1.42%)",
     "Diversified, high-value group ecosystem",
     "Kotak 811 digital-acquisition engine",
     "Consistent high profitability (ROA ~2%)"],
    ["Smallest balance sheet among the five",
     "Lower branch density limits deposit reach",
     "Past RBI restriction on digital onboarding",
     "Founder-to-professional leadership transition",
     "Conservative growth caps market-share gains"],
    ["Scale-up of Kotak 811 & digital lending",
     "Wealth management for HNW customers",
     "Group cross-sell (broking, AMC, insurance)",
     "Selective acquisitions to add scale",
     "Affluent & SME-banking expansion"],
    ["Competition for low-cost deposits",
     "Margin compression in rate-cut cycle",
     "Tech-resilience & regulatory-compliance risk",
     "Aggressive fintech & neo-bank competition",
     "Execution risk in scaling growth"],
)
hr()
page_break()

# ============================ 5. COMPARATIVE STUDY ============================
h1("5. Comparative Study")

h2("5.1 Total Income Comparison (FY22-FY25, Rs Crore, Standalone)")
make_table(
    ["Company", "FY2022", "FY2023", "FY2024", "FY2025"],
    [
        ["SBI", "~3,68,000", "~4,06,973", "~4,66,000", "~5,25,000"],
        ["HDFC Bank", "~1,57,263", "~2,04,666", "~4,07,995", "~4,70,916"],
        ["ICICI Bank", "~1,04,892", "~1,32,716", "~1,71,000", "~2,00,000"],
        ["Axis Bank", "~86,114", "~1,06,155", "~1,37,989", "~1,57,000"],
        ["Kotak Mahindra Bank", "~59,051", "~68,142", "~94,274", "~1,06,902"],
    ],
    widths=[1.8, 1.2, 1.2, 1.2, 1.2],
)
body("SBI remains the largest by total income throughout the period, reflecting its unmatched scale. HDFC Bank's total income roughly doubled between FY23 and FY24 owing to the merger with HDFC Ltd. ICICI Bank, Axis Bank, and Kotak Mahindra Bank all recorded steady double-digit income growth, driven by loan-book expansion and higher interest income in a rising-rate environment. Figures are approximate and on a standalone basis.")

h2("5.2 Financial Performance Comparison (FY25)")
make_table(
    ["Metric", "SBI", "HDFC Bank", "ICICI Bank", "Axis Bank", "Kotak Bank"],
    [
        ["Net Profit (Rs cr)", "~70,901", "~67,347", "~47,227", "~26,373", "~16,450"],
        ["Net Interest Margin", "~3.0%", "~3.5%", "~4.4%", "~3.98%", "~5.0%"],
        ["Return on Assets", "~1.10%", "~1.8%", "~2.4%", "~1.77%", "~2.0%"],
        ["Return on Equity", "~19-20%", "~14%", "~18%", "~16.9%", "~13-14%"],
        ["Gross NPA", "~1.82%", "~1.33%", "~1.67%", "~1.28%", "~1.42%"],
        ["Net NPA", "~0.47%", "~0.43%", "~0.39%", "~0.33%", "~0.31%"],
        ["Capital Adequacy", "~14%", "~19.6%", "~16%+", "~17%", "~22%+"],
    ],
    widths=[1.5, 1.0, 1.1, 1.1, 1.0, 1.0],
)
body("ICICI Bank leads on return on assets and Kotak on net interest margin, reflecting their profitability-first models. SBI leads on absolute net profit and return on equity, driven by scale and improved efficiency. HDFC Bank and Kotak have the strongest capital adequacy, while net NPAs are tightly clustered below 0.5% across all five, signalling system-wide asset-quality strength.")

h2("5.3 Marketing Comparison")
make_table(
    ["Parameter", "SBI", "HDFC Bank", "ICICI Bank", "Axis Bank", "Kotak Bank"],
    [
        ["Brand Strength", "*****", "*****", "****", "***", "****"],
        ["Digital Marketing", "****", "****", "*****", "****", "****"],
        ["Flagship Platform", "YONO", "PayZapp/MobileApp", "iMobile Pay", "open by Axis", "Kotak 811"],
        ["Customer Reach", "*****", "*****", "****", "****", "***"],
        ["Premium Positioning", "***", "****", "****", "****", "*****"],
    ],
    widths=[1.5, 1.0, 1.3, 1.1, 1.1, 1.0],
)

h2("5.4 Operational Comparison (FY25)")
make_table(
    ["Parameter", "SBI", "HDFC Bank", "ICICI Bank", "Axis Bank", "Kotak Bank"],
    [
        ["Net Profit (FY25)", "~Rs 70,901 cr", "~Rs 67,347 cr", "~Rs 47,227 cr", "~Rs 26,373 cr", "~Rs 16,450 cr"],
        ["Employees", "~2,36,000", "~2,14,000", "~1,33,000", "~1,05,000", "~74,000"],
        ["Ownership", "Government-majority", "Private", "Private", "Private", "Private"],
        ["Digital Platform", "YONO", "Digital Factory", "iMobile/InstaBIZ", "open by Axis", "Kotak 811"],
        ["CASA Ratio", "~40%", "~38%", "~39%", "~41%", "~43%"],
    ],
    widths=[1.4, 1.2, 1.2, 1.1, 1.1, 1.0],
)

h2("5.5 HR Comparison")
make_table(
    ["Parameter", "SBI", "HDFC Bank", "ICICI Bank", "Axis Bank", "Kotak Bank"],
    [
        ["Workforce Size", "~2,36,000", "~2,14,000", "~1,33,000", "~1,05,000", "~74,000"],
        ["Recruitment Model", "National exams", "Campus + lateral", "Campus + lateral", "Campus + lateral", "Campus + lateral"],
        ["Culture USP", "Service/Trust", "Performance/Scale", "Risk discipline", "Customer obsession", "Prudence"],
        ["Employee Strength", "Highest", "Very High", "High", "High", "Moderate"],
    ],
    widths=[1.4, 1.1, 1.2, 1.2, 1.2, 1.0],
)
hr()
page_break()
print("sections 4-5 loaded")


# ============================ 6. CHALLENGES & RISK ANALYSIS ============================
h1("6. Challenges & Risk Analysis")
h2("6.1 Current Challenges for Each Company")
h3("HDFC Bank")
body("Following its merger with HDFC Ltd, HDFC Bank inherited a high credit-deposit ratio and a large pool of borrowings, requiring it to mobilise deposits aggressively to normalise its funding profile. This has temporarily moderated loan growth and pressured net interest margins. Integrating the systems, people, and culture of the merged entity while maintaining asset quality and service standards remains a key execution challenge.")
h3("ICICI Bank")
body("ICICI Bank must sustain its sector-leading profitability amid intensifying competition for deposits and a likely compression of margins in a falling-rate environment. Managing asset quality in its growing unsecured retail and business-banking books through the credit cycle, and continuing to improve productivity while rationalising headcount, are its principal near-term challenges.")
h3("State Bank of India")
body("SBI's structurally lower net interest margin and return on assets relative to leading private banks reflect its scale, social obligations, and government-banking mandate. Protecting margins in a rate-cut cycle, managing asset quality across its large agriculture and SME portfolios, modernising legacy technology at enormous scale, and periodically raising capital to fund growth are ongoing challenges.")
h3("Axis Bank")
body("Axis Bank's lower CASA ratio relative to HDFC Bank and ICICI Bank raises its cost of funds and pressures margins. Successfully integrating the acquired Citibank consumer-banking portfolio, deepening its low-cost deposit franchise, and sustaining the recent improvement in asset quality through the unsecured-lending cycle are its central priorities.")
h3("Kotak Mahindra Bank")
body("Kotak's conservative, profitability-first approach has historically capped its market-share gains relative to faster-growing peers. The RBI's earlier restriction on onboarding new customers digitally and issuing new credit cards highlighted technology-resilience gaps that the bank has been remediating. Scaling growth and digital capability while preserving its prudent risk culture, and completing the founder-to-professional leadership transition, are its key challenges.")
h2("6.2 Common External Risks Across the Sector")
for r in [
    "Net interest margin compression: An anticipated interest-rate-cut cycle by the RBI, combined with intense competition for deposits, could squeeze margins across all five banks as lending yields reprice faster than deposit costs.",
    "Deposit-mobilisation pressure: Slower deposit growth relative to credit growth has tightened system liquidity, raising the cost and competition for low-cost CASA and retail term deposits.",
    "Asset-quality risk in unsecured lending: Rapid growth in unsecured personal loans, credit cards, and microfinance has prompted regulatory caution; a turn in the credit cycle could raise slippages.",
    "Regulatory and compliance burden: Evolving RBI norms on capital, liquidity, digital lending, governance, and customer protection, along with the Digital Personal Data Protection Act, 2023, raise compliance costs.",
    "Cybersecurity and technology resilience: Rising digital-transaction volumes increase exposure to cyberattacks, fraud, and outages, with regulators demanding robust resilience and business continuity.",
    "Fintech and Big Tech disruption: Fintechs, neo-banks, and Big Tech players continue to disaggregate payments and lending, pressuring fee income and customer ownership.",
    "Macroeconomic and global risks: Global growth uncertainty, geopolitical tensions, and commodity-price volatility can affect credit demand, asset quality, and treasury income.",
]:
    bullet(r)
hr()
page_break()

# ============================ 7. FUTURE GROWTH OPPORTUNITIES ============================
h1("7. Future Growth Opportunities")
h3("Industry Trends")
body("The Indian financial-services sector is being reshaped by structural trends: (1) the continued explosion of digital payments led by UPI, which is expanding into credit, cross-border, and offline use cases; (2) the rise of account aggregators and cash-flow-based lending, enabling data-driven credit to thin-file customers and MSMEs; (3) co-lending and embedded finance, blending the reach of banks with the agility of fintechs and NBFCs; (4) the shift of household savings toward market-linked products, boosting wealth management, mutual funds, and insurance; and (5) the deployment of AI and generative AI across underwriting, customer service, and fraud management.")
h3("Company Readiness for Growth")
body("SBI is best positioned to capture mass-market and rural growth through its unmatched reach and the YONO platform. HDFC Bank is poised to cross-sell mortgages and the full product suite to its vast deposit base post-merger. ICICI Bank's digital and ecosystem capabilities position it to lead in profitable retail and business-banking growth. Axis Bank is leveraging the Citi acquisition to scale its affluent and cards franchise, while Kotak Mahindra Bank is best placed in high-value wealth management and integrated group cross-sell.")
h3("Expansion Strategies")
body("All five banks are pursuing growth through a mix of organic expansion and selective inorganic moves. SBI is deepening rural and MSME penetration and unlocking value in its subsidiaries. HDFC Bank is expanding semi-urban and rural distribution to mobilise deposits. ICICI and Axis are scaling business banking and affluent segments, with Axis having acquired Citi's India consumer business. Kotak continues to evaluate selective acquisitions to add scale while protecting its conservative risk profile. All are investing heavily in digital platforms and co-lending partnerships.")
h3("Technologies Shaping the Future")
for t in [
    "Generative AI and AI agents for customer service, underwriting, and operations automation",
    "UPI-led innovation: credit on UPI, UPI Lite, and cross-border payment corridors",
    "Account aggregators and cash-flow-based lending for MSMEs and new-to-credit customers",
    "Cloud-native core banking and API-led open banking for embedded finance",
    "Advanced analytics for hyper-personalisation, collections, and real-time fraud detection",
    "Cybersecurity, tokenisation, and resilient infrastructure for digital trust",
]:
    bullet(t)
hr()
page_break()

# ============================ 8. KEY FINDINGS & CONCLUSION ============================
h1("8. Key Findings & Conclusion")
h2("8.1 Best-Performing Company")
body("On a composite basis, SBI and HDFC Bank lead the group. SBI is the best performer by scale and absolute profitability, delivering a record net profit of approximately Rs 70,901 crore in FY25 alongside its strongest-ever asset quality. HDFC Bank leads the private-sector field by balance-sheet size, asset quality, and capital strength. For risk-adjusted profitability, however, ICICI Bank stands out with the highest return on assets in the group.")
h2("8.2 Most Innovative Company")
body("ICICI Bank earns the distinction of being the most innovation-driven of the five, evidenced by its industry-leading iMobile Pay and InstaBIZ apps, its open-architecture API banking platform, and its data-led, 360-degree customer-ecosystem strategy. SBI's YONO platform — one of the largest digital-banking ecosystems in the world by registered users — makes it a close second in terms of digital scale and reach.")
h2("8.3 Strongest Financial Company")
body("ICICI Bank and Kotak Mahindra Bank lead on financial robustness. ICICI delivers the highest return on assets (~2.4%) with strong asset quality, while Kotak combines the highest net interest margin (~5%) with the strongest capital adequacy and the lowest net NPA. SBI ranks highest on absolute profit and return on equity, reflecting its scale advantage.")
h2("8.4 Strongest Marketing Company")
body("HDFC Bank and SBI command the strongest brand equity and customer reach. HDFC Bank leads in cards, payments, and premium positioning, while SBI enjoys unmatched brand trust and national reach built over two centuries. ICICI Bank leads in digital-marketing sophistication, and Kotak holds the strongest premium-banking and wealth-management positioning.")
h2("8.5 Future Industry Outlook")
body("The Indian banking sector's long-term fundamentals are strongly positive. Driven by credit penetration, digitisation, financial inclusion, and the formalisation of the economy, system credit is expected to continue compounding in double digits. While net interest margins face near-term pressure from a likely rate-cut cycle and deposit competition, the sector's well-capitalised balance sheets, multi-year-low NPAs, and record profitability provide a strong foundation. Banks that master digital distribution, low-cost deposit mobilisation, disciplined risk management, and AI-led efficiency will gain durable competitive advantage. All five institutions are well-positioned to remain at the core of India's financial system for the foreseeable future.")
hr()
page_break()

# ============================ 9. RECOMMENDATIONS ============================
h1("9. Recommendations")
h2("9.1 Strategic Recommendations")
for r in [
    "HDFC Bank should prioritise aggressive low-cost deposit mobilisation in semi-urban and rural India to normalise its post-merger credit-deposit ratio and protect margins.",
    "ICICI Bank should continue scaling its high-return Business Banking and SME franchise while maintaining its disciplined risk culture through the credit cycle.",
    "SBI should accelerate the monetisation and value-unlocking of its subsidiaries (SBI Life, SBI Cards, SBI MF) and lean further on YONO to drive cost-efficient growth.",
    "Axis Bank should focus on deepening its CASA franchise and fully integrating the Citi consumer-banking portfolio to strengthen its affluent and cards businesses.",
    "Kotak Mahindra Bank should scale its growth and digital capabilities through Kotak 811 and selective acquisitions while preserving its conservative risk profile.",
]:
    bullet(r)
h2("9.2 Marketing Improvements")
for r in [
    "All five banks should invest more in financial-literacy and trust-building content targeting first-time and digitally native customers.",
    "ICICI and Axis should leverage their premium credit-card and affluent franchises with lifestyle and ecosystem-led co-branded campaigns.",
    "SBI should continue modernising its brand to engage younger customers through YONO while retaining its mass-market trust positioning.",
]:
    bullet(r)
h2("9.3 Financial Suggestions")
for r in [
    "Banks should protect net interest margins through disciplined deposit pricing, optimal loan-mix management, and growth in high-yield retail and SME segments.",
    "SBI should maintain its return-on-assets above 1% through cost efficiency and subsidiary value contribution, supported by adequate capital buffers.",
    "All banks should sustain conservative provisioning and high provision-coverage ratios to absorb any turn in the unsecured-lending credit cycle.",
]:
    bullet(r)
h2("9.4 Operational Improvements")
for r in [
    "All five banks should accelerate digitisation of lending, onboarding, and servicing to lower cost-to-income ratios and improve customer experience.",
    "SBI and HDFC Bank should leverage their scale to drive operating leverage through automation and branch-network optimisation.",
    "Kotak should complete its core-systems modernisation and strengthen technology resilience to support scalable digital growth.",
]:
    bullet(r)
h2("9.5 Technology Recommendations")
for r in [
    "All five banks should establish dedicated generative-AI and AI-agent practices for customer service, underwriting, and operations automation.",
    "Cybersecurity and technology resilience should be elevated to top board-level priorities, with continuous investment in fraud prevention and business continuity.",
    "Banks should build cloud-native, API-led open-banking platforms to power embedded finance, co-lending, and account-aggregator-based credit.",
]:
    bullet(r)
hr()
page_break()

# ============================ 10. REFERENCES ============================
h1("10. References")
h3("Annual Reports & Investor Relations")
refs1 = [
    "State Bank of India - Annual Report FY2024-25 and Q4 FY2025 Results. https://sbi.bank.in/web/investor-relations/annual-report",
    "HDFC Bank - Financial Results (Indian GAAP) for the year ended March 31, 2025. https://www.hdfcbank.com/personal/about-us/investor-relations/financial-results",
    "ICICI Bank - Annual Report 2024-25 and Board's Report. https://www.icicibank.com/ms/aboutus/annual-reports/2024-25/html/board-report.html",
    "Axis Bank - Quarterly & Annual Results FY2024-25. https://www.axisbank.com/quarterly-results/2024-2025/q4/index.html",
    "Kotak Mahindra Bank - Annual Results FY2024-25 and Directors' Report. https://www.indiainfoline.com/company/kotak-mahindra-bank-ltd/reports/directors-report",
]
for i, r in enumerate(refs1, 1):
    doc.add_paragraph(f"{i}. {r}").paragraph_format.space_after = Pt(4)

h3("Industry Reports & Regulatory Sources")
refs2 = [
    "Reserve Bank of India - Report on Trend and Progress of Banking in India and Financial Stability Report. https://www.rbi.org.in",
    "RBI - Database on Indian Economy (deposits, credit, asset-quality data). https://www.rbi.org.in",
    "IBEF - Banking & Financial Services Sector in India. https://www.ibef.org/industry/banking-india",
    "Digital Personal Data Protection Act, 2023 - Government of India. https://www.meity.gov.in/content/digital-personal-data-protection-act-2023",
]
for i, r in enumerate(refs2, 6):
    doc.add_paragraph(f"{i}. {r}").paragraph_format.space_after = Pt(4)

h3("Financial Data & Market Sources")
refs3 = [
    "Quartr - Q4 FY24-25 earnings summaries for SBI, HDFC Bank, ICICI Bank, Axis Bank, and Kotak Mahindra Bank. https://quartr.com",
    "Moneycontrol - Bank financial statements and ratios. https://www.moneycontrol.com",
    "Groww - Bank yearly results and financials. https://groww.in/stocks",
    "ET Money - Kotak Mahindra Bank financials and ratio analysis. https://www.etmoney.com/stocks/kotak-mahindra-bank-ltd/financials",
    "Times of India / Economic Times - FY25 results coverage for SBI and peers. https://timesofindia.indiatimes.com",
]
for i, r in enumerate(refs3, 10):
    doc.add_paragraph(f"{i}. {r}").paragraph_format.space_after = Pt(4)

doc.add_paragraph()
center("- End of Report -", 12, bold=True, color=GREY)

# ============================ SAVE ============================
out = "/projects/sandbox/Hatch/work/Financial_Services_Comparative_Report.docx"
doc.save(out)
print("SAVED:", out)
