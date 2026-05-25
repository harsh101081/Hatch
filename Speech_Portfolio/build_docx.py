"""
Builds the Independent Speech & Public Speaking Portfolio Word document.
Theme: Navy (#0B2545) + Gold (#C9A227) + Off-white (#F5F7FA) + Charcoal (#1F2937)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- THEME ----------
NAVY = RGBColor(0x0B, 0x25, 0x45)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
CHARCOAL = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GREY = RGBColor(0x9C, 0xA3, 0xAF)

HEADING_FONT = "Georgia"   # Playfair Display fallback
BODY_FONT = "Calibri"      # Inter fallback


# ---------- HELPERS ----------
def set_run(run, *, font=BODY_FONT, size=11, bold=False, italic=False, color=CHARCOAL):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)


def add_para(doc, text, *, font=BODY_FONT, size=11, bold=False, italic=False,
             color=CHARCOAL, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        set_run(run, font=font, size=size, bold=bold, italic=italic, color=color)
    return p


def add_h1(doc, text):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    set_run(run, font=HEADING_FONT, size=22, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(2)
    add_gold_rule(doc)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, font=HEADING_FONT, size=14, bold=True, color=NAVY)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, font=BODY_FONT, size=12, bold=True, color=GOLD)


def add_gold_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9A227')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullet(doc, text, *, indent=0.25, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r1 = p.add_run(bold_lead + " ")
        set_run(r1, size=11, bold=True, color=NAVY)
        r2 = p.add_run(text)
        set_run(r2, size=11, color=CHARCOAL)
    else:
        run = p.add_run(text)
        set_run(run, size=11, color=CHARCOAL)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'C9A227')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_run(run, size=11, italic=True, color=CHARCOAL)


def add_script_para(doc, text):
    """For the speech script — bigger, more readable line spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.45
    # Process **bold** and *italic* and (pause) in italics
    import re
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|\([^)]+\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            set_run(r, size=11.5, bold=True, color=NAVY)
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1])
            set_run(r, size=11.5, italic=True, color=GREY)
        elif part.startswith('(') and part.endswith(')'):
            r = p.add_run(part)
            set_run(r, size=10.5, italic=True, color=GOLD)
        else:
            r = p.add_run(part)
            set_run(r, size=11.5, color=CHARCOAL)
    return p


def add_table(doc, headers, rows, *, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run(run, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # navy fill
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0B2545')
        tcPr.append(shd)
    # Body
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            set_run(run, size=10.5, color=CHARCOAL)
            if r_idx % 2 == 1:
                tcPr = cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F7FA')
                tcPr.append(shd)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'D1D5DB')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    return table


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ---------- DOCUMENT ----------
doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Default style
style = doc.styles['Normal']
style.font.name = BODY_FONT
style.font.size = Pt(11)
style.font.color.rgb = CHARCOAL


# ============== COVER PAGE ==============
for _ in range(4):
    doc.add_paragraph()

# Eyebrow
add_para(doc, "BACHELOR OF FINANCIAL MARKETS  |  2025–2026",
         size=10, bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# Big title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Independent Speech &")
set_run(r, font=HEADING_FONT, size=30, bold=True, color=NAVY)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Public Speaking Portfolio")
set_run(r2, font=HEADING_FONT, size=30, bold=True, color=NAVY)
p2.paragraph_format.space_after = Pt(14)

# Gold rule
add_gold_rule(doc)

add_para(doc,
         '"Why Financial Literacy Is the Most Important Skill for My Generation"',
         size=14, italic=True, color=CHARCOAL,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

add_para(doc, "A Speech, Slide Deck & Reflection Portfolio for Internship Readiness",
         size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

# Submission box
info = [
    ("Submitted by", "[Student Name]"),
    ("Roll Number", "[____________]"),
    ("Course", "Bachelor of Financial Markets (BFM)"),
    ("Year / Semester", "[____________]"),
    ("College", "[____________]"),
    ("Submitted to", "[Faculty Name]"),
    ("Date", "[____________]"),
]

t = doc.add_table(rows=len(info), cols=2)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, (k, v) in enumerate(info):
    c1, c2 = t.rows[i].cells
    c1.text = ""
    c2.text = ""
    r1 = c1.paragraphs[0].add_run(k)
    set_run(r1, size=10.5, bold=True, color=NAVY)
    r2 = c2.paragraphs[0].add_run(v)
    set_run(r2, size=10.5, color=CHARCOAL)
    c1.width = Inches(2.0)
    c2.width = Inches(3.5)

page_break(doc)


# ============== TABLE OF CONTENTS ==============
add_h1(doc, "Table of Contents")

toc_rows = [
    ("1", "Cover Page"),
    ("2", "Topic & Objective"),
    ("3", "Full Speech Script (5–7 minutes)"),
    ("4", "Speech Outline & Cue Cards"),
    ("5", "Slide Deck Content (8 Slides)"),
    ("6", "Video Recording Guide"),
    ("7", "Self-Evaluation"),
    ("8", "Public Speaking Portfolio"),
    ("9", "Reflection"),
    ("10", "Final Formatting Guide"),
    ("11", "Final Output Structure"),
]
add_table(doc, ["#", "Section"], toc_rows, col_widths=[0.7, 5.3])
page_break(doc)


# ============== SECTION 2 — TOPIC & OBJECTIVE ==============
add_h1(doc, "Section 2 — Topic & Objective")

add_h2(doc, "Final Speech Topic")
add_quote(doc, '"Why Financial Literacy Is the Most Important Skill for My Generation."')

add_h2(doc, "Why This Topic Was Selected")
add_para(doc,
    "As a BFM student, I am surrounded by markets, investing, and financial decisions every day. "
    "But outside the classroom, I see a different reality. Most students my age know how to use ten "
    "different apps but cannot explain what an SIP, an index fund, or a credit score actually is. "
    "I chose this topic because it sits exactly between what I am studying and what my generation "
    "actually needs to hear. It is also a topic I can speak about honestly — not as an expert, but "
    "as a student who is learning the same things I am asking my audience to learn.")

add_h2(doc, "Objective of the Speech")
add_bullet(doc, "explain what financial literacy means in simple, relatable language.", bold_lead="To")
add_bullet(doc, "highlight why young Indians, especially students, must learn it early.", bold_lead="To")
add_bullet(doc, "inspire the audience to take one small financial action after listening.", bold_lead="To")

add_h2(doc, "Intended Audience")
add_bullet(doc, "College students (16–22 years old)")
add_bullet(doc, "BFM, BBA, B.Com, and general degree students")
add_bullet(doc, "Faculty and internship interviewers reviewing this portfolio")

add_h2(doc, "What the Audience Should Learn")
add_bullet(doc, "Financial literacy is not 'advanced finance' — it is a life skill.")
add_bullet(doc, "Starting early is more powerful than starting big.")
add_bullet(doc, "Three habits — Track, Save, Invest — are enough to begin with.")
add_bullet(doc, "The best time to learn money is before you start earning it.")

page_break(doc)


# ============== SECTION 3 — FULL SPEECH SCRIPT ==============
add_h1(doc, "Section 3 — Full Speech Script")

add_para(doc,
    "Length: ~870 spoken words → roughly 6 minutes at a calm, confident pace.",
    size=10, italic=True, color=GREY, space_after=2)
add_para(doc,
    "Notation: (pause) = short breath, **bold** = emphasis, *italics* = softer/personal tone.",
    size=10, italic=True, color=GREY, space_after=10)

add_h3(doc, "INTRODUCTION — The Hook (≈ 60 seconds)")
add_script_para(doc, "Good morning, everyone. (pause)")
add_script_para(doc, "Before I begin, I want to ask you a quick question. Just answer it in your head — you don't have to say it out loud.")
add_script_para(doc, "**How much money did you spend last month?** (pause)")
add_script_para(doc, "Now a harder one — **where did it actually go?** (pause)")
add_script_para(doc, "If your honest answer is *I have no idea,* don't worry. You are not alone. In fact, you are completely normal. (pause)")
add_script_para(doc, "A survey by the National Centre for Financial Education found that only about **27% of Indians are financially literate.** That means roughly three out of every four people around you — your friends, your cousins, even some adults you look up to — are making money decisions without really understanding money.")
add_script_para(doc, "And that is exactly why I am standing here today. (pause)")
add_script_para(doc, "My name is [Your Name], I am a Bachelor of Financial Markets student, and I want to talk to you about something I genuinely believe is **the most important skill of our generation** — *financial literacy.*")

add_h3(doc, "TRANSITION")
add_script_para(doc, "So let me ask the obvious question first — *what does financial literacy even mean?*")

add_h3(doc, "MAIN POINT 1 — What Financial Literacy Actually Is (≈ 75 seconds)")
add_script_para(doc, "When most people hear the word **finance,** they imagine the stock market, traders shouting on TV, or some uncle at a family function giving unsolicited advice about gold. (pause for light laughter)")
add_script_para(doc, "But financial literacy is **much simpler than that.**")
add_script_para(doc, "In one line — financial literacy is the ability to **understand, manage, and grow your money confidently.**")
add_script_para(doc, "That's it.")
add_script_para(doc, "It is knowing the difference between a *need* and a *want.* It is understanding what an **SIP**, a **credit score**, or an **emergency fund** actually does. It is being able to read your own bank statement without panicking.")
add_script_para(doc, "It is **not** about becoming Warren Buffett. (pause) It is about not becoming the person who reaches the end of the month and asks, *Where did my salary go?*")
add_script_para(doc, "And here is the part that surprised me when I started studying BFM — **none of this is taught in school.** We learn trigonometry, we learn the periodic table, but we don't learn what an EMI is until we are signing one.")

add_h3(doc, "TRANSITION")
add_script_para(doc, "Which brings me to my second point — *why this matters so much for people our age.*")

add_h3(doc, "MAIN POINT 2 — Why It Matters Right Now (≈ 90 seconds)")
add_script_para(doc, "We are the **first fully digital money generation** in India. (pause)")
add_script_para(doc, "Think about it. Our parents handled cash. We handle UPI. We send money in two seconds, we shop at midnight, we invest from our phones in three taps. Our money moves **faster than our understanding of it.**")
add_script_para(doc, "And that is dangerous.")
add_script_para(doc, "Because the same phone that lets you start an SIP also lets you lose ten thousand rupees on a 'guaranteed tip' from a random Telegram group. (pause) The same app that helps you save can push you into a loan you don't need.")
add_script_para(doc, "Add one more thing to this — **inflation.** The samosa that cost ten rupees when we were kids costs twenty-five today. (pause) If your money is just sitting in a savings account, it is **slowly losing value every single year.**")
add_script_para(doc, "So the question is not *should I learn about money?* The question is *can I afford not to?*")
add_script_para(doc, "And here is the most beautiful part — **time is on our side.** If a 19-year-old invests just two thousand rupees a month at a reasonable return, by the time they are 45, that small habit can grow into something genuinely life-changing. (pause)")
add_script_para(doc, "Not because of magic. Because of **compounding** — and compounding rewards the people who **start early,** not the people who start big.")

add_h3(doc, "TRANSITION")
add_script_para(doc, "So if you're sitting there thinking, *Okay, I get it, but where do I even start?* — let me make it really simple.")

add_h3(doc, "MAIN POINT 3 — Three Habits To Start With (≈ 75 seconds)")
add_script_para(doc, "You don't need a finance degree. You don't need a Demat account on day one. You only need **three habits.** I call them **Track, Save, Invest.** (emphasis)")
add_script_para(doc, "**One — Track.** For 30 days, just write down every rupee you spend. No judgment, no budget yet. Just awareness. You will be shocked how much disappears into food delivery and 'small' online orders.")
add_script_para(doc, "**Two — Save.** Follow a simple rule — the **50-30-20 rule.** 50% on needs, 30% on wants, 20% on savings. Even if you are a student with pocket money, save **something.** The amount is not the point. The habit is.")
add_script_para(doc, "**Three — Invest.** Once you have a small emergency fund, start an **SIP of even five hundred rupees** in an index fund. Read about it before you do it. Make mistakes with small money now, so you don't make them with big money later.")
add_script_para(doc, "That's it. (pause) **Track. Save. Invest.** Three words. One life skill.")

add_h3(doc, "CONCLUSION — The Closing (≈ 60 seconds)")
add_script_para(doc, "Let me bring this back to where we started. (pause)")
add_script_para(doc, "I asked you what you spent last month and where it went. By next month — I hope at least a few of you can actually answer that question.")
add_script_para(doc, "Because financial literacy is **not** about getting rich quickly. It is about **getting in control quietly.** (pause for emphasis)")
add_script_para(doc, "It is the difference between money working for you, and you working only for money.")
add_script_para(doc, "We are young. We don't have huge salaries yet. But we have something more powerful than any salary — **time, curiosity, and the internet.** And that is more than enough to start.")
add_script_para(doc, "So when you walk out of this room today, I am not asking you to open a brokerage account or read a finance book tonight. I am only asking you to do **one thing** — *track every rupee you spend for the next seven days.*")
add_script_para(doc, "That one small habit is where my financial journey started. (pause) And I genuinely believe it is where yours can begin too.")
add_script_para(doc, "Thank you. (pause, smile, eye contact)")

page_break(doc)


# ============== SECTION 4 — OUTLINE & CUE CARDS ==============
add_h1(doc, "Section 4 — Speech Outline & Cue Cards")

add_h2(doc, "A) Short Outline (One-Page Version)")

add_h3(doc, "1. Introduction")
add_bullet(doc, 'Hook: "How much did you spend last month?"')
add_bullet(doc, "Stat: Only 27% of Indians are financially literate.")
add_bullet(doc, "Self-introduction + topic announcement.")

add_h3(doc, "2. Main Point 1 — What Financial Literacy Means")
add_bullet(doc, "Not stock-market jargon — it's a life skill.")
add_bullet(doc, "Definition: understand, manage, grow money confidently.")
add_bullet(doc, "Why it's missing from school education.")

add_h3(doc, "3. Main Point 2 — Why It Matters Now")
add_bullet(doc, "First fully digital money generation.")
add_bullet(doc, "Risks: scams, impulsive loans, inflation.")
add_bullet(doc, "Power of starting early — compounding.")

add_h3(doc, "4. Main Point 3 — Three Habits: Track, Save, Invest")
add_bullet(doc, "Track: 30 days of expenses.")
add_bullet(doc, "Save: 50-30-20 rule.")
add_bullet(doc, "Invest: ₹500 SIP in an index fund.")

add_h3(doc, "5. Conclusion")
add_bullet(doc, "Callback to opening question.")
add_bullet(doc, '"Money working for you vs. you working only for money."')
add_bullet(doc, "Call to action: track every rupee for 7 days.")
add_bullet(doc, "Thank you.")

add_h2(doc, "B) Cue Cards (Speaker Notes)")
add_para(doc, "Use one cue card per section. Keep cards small. Eye contact > reading.",
         size=10, italic=True, color=GREY, space_after=8)

cards = [
    ("CARD 1 — INTRO", [
        "Smile. Pause. Eye contact across room.",
        'Ask: "How much did you spend last month?"',
        "Drop stat: 27% financially literate.",
        "Introduce yourself + BFM + topic."
    ]),
    ("CARD 2 — DEFINITION", [
        "Bust the myth: finance ≠ stock screens.",
        "Define in one line.",
        '"Not about being Buffett — about not asking where did my salary go?"',
        "Mention: not taught in school."
    ]),
    ("CARD 3 — WHY NOW", [
        '"First digital-money generation."',
        "UPI, instant loans, scams, Telegram tips.",
        "Inflation example: ₹10 samosa → ₹25.",
        "Compounding rewards early, not big."
    ]),
    ("CARD 4 — THREE HABITS", [
        "TRACK → 30 days, no judgment.",
        "SAVE → 50-30-20.",
        "INVEST → ₹500 SIP, index fund.",
        'Chant: "Track. Save. Invest."'
    ]),
    ("CARD 5 — CLOSE", [
        "Callback to opening.",
        '"Control quietly, not rich quickly."',
        "CTA: track every rupee for 7 days.",
        "Thank you. Hold eye contact."
    ]),
]
for title, items in cards:
    add_h3(doc, title)
    for it in items:
        add_bullet(doc, it)

add_h2(doc, "C) Cue Words Strip (Presenter View)")
add_quote(doc, "HOOK → STAT → ME → DEFINE → SCHOOL-GAP → DIGITAL GEN → INFLATION → COMPOUND → TRACK → SAVE → INVEST → CALLBACK → CTA → THANK YOU")

page_break(doc)


# ============== SECTION 5 — SLIDE DECK CONTENT ==============
add_h1(doc, "Section 5 — Slide Deck Content")

add_para(doc, "Total slides: 8 | Theme: Navy #0B2545 + Gold #C9A227 + Off-white #F5F7FA",
         size=10, italic=True, color=GREY, space_after=10)

slides = [
    {
        "n": 1, "title": "Title Slide",
        "on": ["Title: Why Financial Literacy Is the Most Important Skill for My Generation",
               "Subtitle: An Independent Speech by [Student Name] | BFM"],
        "visual": "Faint background line graph trending upward; small gold circle accent.",
        "icon": "Minimal upward arrow or candle chart in gold.",
        "layout": "Centered text, large title, small subtitle below a thin gold rule.",
        "notes": "Pause for 2 seconds. Smile. Then begin with the hook question."
    },
    {
        "n": 2, "title": "The Hook",
        "on": ["How much did you spend last month?",
               "Where did it actually go?"],
        "visual": "A blurred wallet or UPI icon faintly in the corner.",
        "icon": "Magnifying glass over a rupee symbol.",
        "layout": "Two bold lines stacked, large type, lots of white space.",
        "notes": "Let the questions sit. Don't rush. Let the audience think."
    },
    {
        "n": 3, "title": "The Problem",
        "on": ["Only 27% of Indians are financially literate.",
               "3 out of 4 people around you don't fully understand their own money."],
        "visual": "Big infographic — 4 person icons, 1 highlighted in gold, 3 in light grey.",
        "icon": "Pie chart (27% gold / 73% grey).",
        "layout": "Stat on the left, infographic on the right.",
        "notes": "Emphasize the number. Pause before the next slide."
    },
    {
        "n": 4, "title": "What It Really Means",
        "on": ["Financial Literacy = A Life Skill",
               "• Understand money",
               "• Manage money",
               "• Grow money — confidently"],
        "visual": "Three small icons in a row (book, wallet, plant growing).",
        "icon": "Brain → Wallet → Sprout.",
        "layout": "Horizontal three-column layout, equal spacing.",
        "notes": "Stress: 'It's not Buffett-level finance. It's daily-life finance.'"
    },
    {
        "n": 5, "title": "Why Now",
        "on": ["Our money moves faster than our understanding.",
               "• UPI, instant loans, app-based investing",
               "• Scams, impulsive spending, inflation",
               "• Time is our biggest advantage"],
        "visual": "Phone graphic with money flying out + a clock icon.",
        "icon": "Smartphone + clock + small upward arrow.",
        "layout": "Bullets on the left, phone illustration on the right.",
        "notes": "Use the samosa example here for warmth and humor."
    },
    {
        "n": 6, "title": "The Power of Starting Early",
        "on": ["Compounding rewards the early, not the big.",
               "₹2,000/month from age 19 → can grow into something life-changing by 45."],
        "visual": "A simple line chart curving upward (exponential), gold line on navy grid.",
        "icon": "Clock + upward graph.",
        "layout": "Stat on top, chart fills lower 60% of slide.",
        "notes": "Don't get technical. Say: 'Start early — that's the entire secret.'"
    },
    {
        "n": 7, "title": "Three Habits. One Life Skill.",
        "on": ["TRACK — every rupee for 30 days",
               "SAVE — 50/30/20 rule",
               "INVEST — start with ₹500 SIP"],
        "visual": "Three stacked cards or tiles, each with one icon.",
        "icon": "Notebook → Piggy bank → Upward chart.",
        "layout": "Three equal columns with bold keywords on top, one-line under each.",
        "notes": "Slow down. Repeat: 'Track. Save. Invest.' Let it land."
    },
    {
        "n": 8, "title": "Conclusion / Close",
        "on": ["It's not about getting rich quickly. It's about getting in control quietly.",
               "Track every rupee you spend for the next 7 days. That's where it begins.",
               "Thank You.",
               "[Student Name] | BFM"],
        "visual": "Soft navy background, gold rule, minimal type.",
        "icon": "Small gold circle or seal with initials.",
        "layout": "Centered, calm, almost like the cover slide.",
        "notes": "Pause. Maintain eye contact. Smile. Wait for applause before walking off."
    },
]

for s in slides:
    add_h2(doc, f"Slide {s['n']} — {s['title']}")
    add_h3(doc, "On-slide text")
    for line in s["on"]:
        add_bullet(doc, line)
    add_h3(doc, "Visual")
    add_para(doc, s["visual"], size=11)
    add_h3(doc, "Icon Idea")
    add_para(doc, s["icon"], size=11)
    add_h3(doc, "Layout")
    add_para(doc, s["layout"], size=11)
    add_h3(doc, "Speaker Notes")
    add_quote(doc, s["notes"])

page_break(doc)


# ============== SECTION 6 — VIDEO RECORDING GUIDE ==============
add_h1(doc, "Section 6 — Video Recording Guide")

guide = [
    ("1. Camera Setup", [
        "Use a phone with at least 1080p video.",
        "Mount it on a tripod, books, or stable surface at eye level.",
        "Frame yourself from mid-chest to just above the head.",
        "Keep the slide screen visible next to you for split-screen later."
    ]),
    ("2. Lighting", [
        "Face a window during the daytime — natural light is the best free upgrade.",
        "At night: ring light or plain white wall + lamp pointed at the ceiling.",
        "Avoid backlight (window behind you = silhouette)."
    ]),
    ("3. Audio", [
        "Use earphones with mic or a clip-on lavalier (₹300–₹600 ones work fine).",
        "Record in a quiet room — close windows, switch off noisy fans.",
        "Do a 10-second test recording before the real take."
    ]),
    ("4. Body Language", [
        "Stand, don't sit. Standing makes your voice more confident.",
        "Shoulders relaxed, feet shoulder-width apart.",
        "Don't lock your knees or sway side to side."
    ]),
    ("5. Hand Gestures", [
        "Use hands at chest level, not below the waist.",
        "Open palms = trust. Pointing fingers = aggressive (avoid).",
        "Match gestures to keywords: 'three habits' → show three fingers."
    ]),
    ("6. Eye Contact", [
        "Look directly into the camera lens, not at your own face on the screen.",
        "Imagine one specific friend on the other side of the lens. Speak to them."
    ]),
    ("7. Speaking Speed", [
        "Aim for ~140–150 words per minute.",
        "Pauses are powerful. Use them after questions and big statements.",
        "If you stumble, don't apologize — pause, breathe, continue."
    ]),
    ("8. Confidence Tips", [
        "Do two warm-up reads before the final take.",
        "Smile slightly before you hit record — it carries into your voice.",
        "Power-pose for 60 seconds before recording. Sounds silly. Works."
    ]),
    ("9. What to Wear", [
        "Solid colors only — navy, white, light blue, beige, soft maroon.",
        "Avoid logos, busy patterns, neon shades, and pure black on dark backgrounds.",
        "A blazer over a plain shirt looks internship-ready without being stiff.",
        "Hair neat, face clean, minimal accessories."
    ]),
    ("10. Common Mistakes to Avoid", [
        "Reading directly from a script in front of the camera.",
        "Speaking in a flat, monotonous tone.",
        "Filling silence with 'um', 'like', 'basically', 'you know'.",
        "Cropping your forehead out of the frame.",
        "Recording vertically — always shoot horizontal (landscape).",
        "Forgetting to switch the phone to Do Not Disturb."
    ]),
]
for title, items in guide:
    add_h2(doc, title)
    for it in items:
        add_bullet(doc, it)

add_h2(doc, "Final Take Checklist")
checklist = [
    "Tripod stable, phone horizontal",
    "Light on face, not behind",
    "Mic plugged in and tested",
    "DND mode on",
    "Slide deck open as backup",
    "Water nearby",
    "Two practice runs done"
]
for it in checklist:
    add_bullet(doc, "☐  " + it)

page_break(doc)


# ============== SECTION 7 — SELF EVALUATION ==============
add_h1(doc, "Section 7 — Self-Evaluation")

add_h2(doc, "Strengths Demonstrated")
add_bullet(doc, "I picked a topic I genuinely care about, which made the speech sound authentic rather than rehearsed.", bold_lead="Topic clarity:")
add_bullet(doc, "The speech followed a clear three-point structure — What, Why Now, How — which made it easy for the audience to follow.", bold_lead="Structure:")
add_bullet(doc, "Using everyday examples (UPI, samosa pricing, ₹500 SIP) helped translate finance into language my audience could relate to.", bold_lead="Real examples:")
add_bullet(doc, "I worked on slowing down at key moments instead of rushing through, especially at the hook and the closing line.", bold_lead="Pacing:")

add_h2(doc, "Areas Needing Improvement")
add_bullet(doc, "still feel slightly stiff in the first minute. I tend to keep my hands too close to my body before I 'warm up'.", bold_lead="Hand gestures")
add_bullet(doc, "during the middle section drifts toward the slide screen instead of staying with the audience.", bold_lead="Eye contact")
add_bullet(doc, "I caught myself using 'basically' and 'actually' more than I'd like. I am consciously working on replacing them with pauses.", bold_lead="Filler words —")
add_bullet(doc, "this is where nervousness shows up the most, and it's the most important part of any speech.", bold_lead="Confidence in the opening 15 seconds —")

add_h2(doc, "Confidence Level")
add_para(doc,
    "On a 1–10 scale, my current confidence rating is 7/10. I am comfortable with the content and "
    "structure, and I no longer freeze mid-sentence. The remaining 3 points come down to physical "
    "presence — voice projection, gestures, and stage stillness — which only improve with repetition.")

add_h2(doc, "Slide Usage Effectiveness")
add_para(doc,
    "The slide deck genuinely supported the speech instead of competing with it. I deliberately kept "
    "slides minimal so that the audience would listen to me, not read past me. The strongest slides "
    "were Slide 3 (the 27% statistic) and Slide 7 (Track / Save / Invest), because they gave the "
    "audience something visual to anchor onto.")

add_h2(doc, "Audience Engagement Reflection")
add_para(doc,
    'The opening question — "How much did you spend last month?" — worked well as a hook because I '
    "could see the audience pause and actually think. The samosa-inflation example got smiles, which "
    "lowered the formality of the room. If I were to do this speech again, I would add one more "
    "direct question in the middle to keep engagement consistent throughout.")

add_h2(doc, "Overall Self-Rating")
add_table(doc, ["Criteria", "Rating (out of 10)"], [
    ["Content quality", "9"],
    ["Structure & flow", "8"],
    ["Delivery & confidence", "7"],
    ["Slide design", "8"],
    ["Audience engagement", "7"],
    ["Time management (within 5–7 min)", "9"],
    ["Overall", "8 / 10"],
], col_widths=[4.0, 2.0])

page_break(doc)


# ============== SECTION 8 — PORTFOLIO ==============
add_h1(doc, "Section 8 — Public Speaking Portfolio")

add_h2(doc, "Speaker Profile")
add_quote(doc,
    "[Student Name] — Bachelor of Financial Markets (BFM) student | Aspiring equity research and "
    "finance intern.")
add_para(doc,
    "A finance student who genuinely enjoys breaking down complex market concepts into language a "
    "non-finance audience can understand. Comfortable speaking on topics across investing, financial "
    "literacy, and market behavior. Currently building public speaking and communication skills "
    "through coursework, college events, and independent speech projects.")

add_h2(doc, "Areas of Interest")
for x in [
    "Stock Market Research",
    "Equity Analysis & Valuation Basics",
    "Trading & Investing Fundamentals",
    "Portfolio Construction & Risk Management",
    "Personal Finance & Financial Literacy",
    "The Intersection of Technology and Finance (UPI, fintech, digital investing)"
]:
    add_bullet(doc, x)

add_h2(doc, "Finance-Related Speaking Strengths")
add_bullet(doc, "I can explain finance concepts to a non-finance audience without sounding patronising or overly technical.", bold_lead="Translation skill:")
add_bullet(doc, "I prefer using real-life micro-examples (samosa inflation, ₹500 SIP, UPI scams) over textbook definitions.", bold_lead="Story-led delivery:")
add_bullet(doc, "Every speech I prepare follows a clear What → Why → How structure.", bold_lead="Structured thinking:")
add_bullet(doc, "I design content around what the listener needs, not what the speaker wants to show off.", bold_lead="Audience-first framing:")

add_h2(doc, "Comfortable Presentation Topics")
add_table(doc, ["#", "Topic", "Length"], [
    ["1", "Why Financial Literacy Is the Most Important Skill for My Generation", "5–7 min"],
    ["2", "Investing for Beginners — Where to Actually Start", "5–7 min"],
    ["3", "Why You Should Research Before You Invest", "5–10 min"],
    ["4", "Understanding Risk Before Chasing Returns", "5–7 min"],
    ["5", "How Technology Is Quietly Rewiring Indian Finance", "7–10 min"],
    ["6", "Building Wealth Early: The Math Behind Starting at 19", "5–7 min"],
], col_widths=[0.5, 4.5, 1.0])

add_h2(doc, "Communication Skills Snapshot")
add_table(doc, ["Skill", "Self-Rating"], [
    ["Verbal communication (English)", "★★★★☆"],
    ["Public speaking & stage presence", "★★★★☆"],
    ["Presentation design (PowerPoint / Keynote)", "★★★★☆"],
    ["Storytelling & audience engagement", "★★★★☆"],
    ["Listening & responding under questions", "★★★☆☆"],
    ["Confidence in interview-style settings", "★★★★☆"],
], col_widths=[4.0, 2.0])

add_h2(doc, "Key Learning Outcomes")
for x in [
    "Learned how to convert a finance topic into a narrative, not a lecture.",
    "Built the discipline of writing, rehearsing, and timing a structured 6-minute speech.",
    "Learned to design slides that serve the speaker, not replace them.",
    "Practised recording, reviewing, and self-critiquing my own delivery.",
    "Improved at handling nervousness through preparation rather than avoidance."
]:
    add_bullet(doc, x)

add_h2(doc, "Public Speaking Achievements & Activities")
add_para(doc, "(To be customised by the student.)", size=10, italic=True, color=GREY)
for x in [
    "Independent Speech Project — Financial Literacy, BFM Coursework, 2025–26",
    "[Add] College finance club presentation / market update",
    "[Add] Group case-study presentation in BFM coursework",
    "[Add] Inter-college fest event participation",
    "[Add] Online certification — communication / public speaking"
]:
    add_bullet(doc, x)

add_h2(doc, "Professional Growth Summary")
add_quote(doc,
    "Through this project, I moved from being someone who 'knew about finance' to someone who can "
    "stand up and explain it. That shift — from understanding to articulation — is the single "
    "biggest reason I now feel internship-ready. Whether it is presenting a stock pitch, walking "
    "an interviewer through my thought process, or speaking at a finance club event, I now trust "
    "myself to communicate ideas with structure, calmness, and clarity.")

page_break(doc)


# ============== SECTION 9 — REFLECTION ==============
add_h1(doc, "Section 9 — Reflection")

add_h2(doc, "Why I Chose This Topic")
add_para(doc,
    "I chose financial literacy because it is the rare topic that sits inside my course and inside "
    "my generation's daily life at the same time. Most of my classmates can name three IPL teams "
    "faster than three mutual funds. I didn't want to give a speech that sounded impressive in a "
    "classroom and useless outside of it. I wanted to talk about something that could actually "
    "change one habit in one listener — and financial literacy felt like that topic.")

add_h2(doc, "Challenges During Preparation")
add_para(doc,
    "The hardest part wasn't writing the speech. It was deciding what to leave out. As a BFM student, "
    "I had the temptation to talk about beta, alpha, equity research, technical analysis, valuation "
    "— all the things that make me sound smart. But every time I added jargon, the speech became "
    "colder. I had to keep cutting until what was left actually sounded like me, not like a textbook.")
add_para(doc,
    "The second challenge was timing. My first draft was over 9 minutes. Cutting it to under 7 "
    "without losing meaning forced me to write tighter sentences — a skill I now realise applies "
    "to emails, interviews, and even research notes.")

add_h2(doc, "The Speaking Experience Itself")
add_para(doc,
    "Standing up to deliver this was honestly nerve-wracking the first few times. My voice was "
    "faster than I wanted it to be, and I could feel my hands looking for something to hold. By the "
    "third rehearsal, the words had moved from my notes into my memory, and that's when I started "
    "actually speaking instead of reciting.")
add_para(doc,
    "The recorded final take felt different — calmer, slower, more eye contact, fewer fillers. "
    "Watching myself back was uncomfortable, but it was also the single most useful learning "
    "moment of this entire project.")

add_h2(doc, "Nervousness vs. Confidence")
add_para(doc,
    "I used to think nervousness meant I was unprepared. Through this project I learned that "
    "nervousness mostly means the topic matters to you. That reframe changed everything. Now, "
    "before I begin, I take three slow breaths and tell myself: 'You're not nervous. You just "
    "care.' It works almost every time.")

add_h2(doc, "Lessons Learned")
for x in [
    "A speech is a conversation, not a performance.",
    "The audience does not care about your slides. They care about whether you sound believable.",
    "The opening 15 seconds and the final 15 seconds are 80% of what people remember.",
    "Pauses are not silence — they are emphasis.",
    "Preparation kills nervousness faster than confidence ever will."
]:
    add_bullet(doc, x)

add_h2(doc, "Future Improvement Goals")
for x in [
    "Speak in at least one finance-related college event per semester.",
    "Sit for a mock equity-pitch presentation and get faculty/peer feedback.",
    "Reduce filler words from 4–5 per minute to under 2 per minute.",
    "Build a small library of 5 strong, ready-to-deliver finance talks.",
    "Eventually, present a research-based pitch on a real listed company."
]:
    add_bullet(doc, x)

page_break(doc)


# ============== SECTION 10 — FORMATTING ==============
add_h1(doc, "Section 10 — Final Formatting Guide")

add_h2(doc, "Recommended Color Palette")
add_table(doc, ["Role", "Color", "Hex"], [
    ["Primary (deep navy)", "Navy Blue", "#0B2545"],
    ["Secondary (background)", "Off-White", "#F5F7FA"],
    ["Accent (premium feel)", "Gold", "#C9A227"],
    ["Body text", "Charcoal", "#1F2937"],
    ["Subtle dividers", "Cool Grey", "#9CA3AF"],
], col_widths=[2.5, 2.0, 1.5])

add_h2(doc, "Recommended Fonts")
add_table(doc, ["Use", "Font", "Backup"], [
    ["Major headings", "Playfair Display", "Merriweather / Georgia"],
    ["Subheadings", "Inter Semi-Bold", "Lato Bold"],
    ["Body text", "Inter Regular", "Calibri / Lato"],
    ["Stat numbers", "Inter Bold", "Montserrat Bold"],
], col_widths=[2.0, 2.0, 2.0])

add_h2(doc, "Layout & Structure")
for x in [
    "Page size: A4 portrait, 1-inch margins.",
    "Line spacing: 1.4 in body, 1.15 in tables.",
    "Section dividers: gold horizontal rule + section number + name.",
    "Headers: section name top-left, page number top-right.",
    "Footers: '[Student Name] | BFM | Independent Speech Portfolio' in 9pt grey.",
    "Page breaks: each major section on a new page.",
    "Cover page: no header/footer; logo at top, gold rule, centered text."
]:
    add_bullet(doc, x)

add_h2(doc, "Visual Aesthetic")
for x in [
    "Theme: modern minimal finance — Bloomberg meets a well-designed annual report.",
    "Iconography: outline icons only (Lucide, Tabler, Phosphor). No 3D clipart.",
    "Imagery: faint background line graphs / candlestick patterns at low opacity.",
    "Treat white space as a design element, not empty space."
]:
    add_bullet(doc, x)

page_break(doc)


# ============== SECTION 11 — FINAL OUTPUT STRUCTURE ==============
add_h1(doc, "Section 11 — Final Output Structure")

add_table(doc, ["Page", "Content"], [
    ["1", "Cover Page — Title, student name, course, submission line"],
    ["2", "Table of Contents"],
    ["3", "Section 2 — Topic & Objective"],
    ["4–6", "Section 3 — Full Speech Script"],
    ["7", "Section 4 — Speech Outline + Cue Cards"],
    ["8–10", "Section 5 — Slide Deck Content"],
    ["11", "Section 6 — Video Recording Guide"],
    ["12", "Section 7 — Self-Evaluation"],
    ["13", "Section 8 — Public Speaking Portfolio"],
    ["14", "Section 9 — Reflection"],
    ["15", "Section 10 — Final Formatting Guide"],
    ["16", "Closing Page — Thank You + Contact"],
], col_widths=[1.0, 5.0])

add_h2(doc, "Export Checklist")
for x in [
    "Replace every [Student Name], [College], [Date] placeholder",
    "Add personal speaking achievements in Section 8",
    "Insert headshot on cover page (optional but recommended)",
    "Export to PDF (primary submission file)",
    "Keep an editable .docx copy for future updates",
    "Build the .pptx version of Section 5 for actual presentation",
    "Embed or link the final recorded video (YouTube unlisted / Google Drive)",
    "Proofread once aloud — not just visually"
]:
    add_bullet(doc, "☐  " + x)

# ---------- CLOSING PAGE ----------
page_break(doc)
for _ in range(6):
    doc.add_paragraph()
add_para(doc, "T H A N K   Y O U", font=HEADING_FONT, size=32, bold=True,
         color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_gold_rule(doc)
add_para(doc, "Independent Speech & Public Speaking Portfolio",
         size=12, italic=True, color=CHARCOAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "Bachelor of Financial Markets (BFM)",
         size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para(doc, "[Student Name]", size=14, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "[Email]  |  [LinkedIn]  |  [College Name]",
         size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------- SAVE ----------
output = "Hatch/Speech_Portfolio/Independent_Speech_Public_Speaking_Portfolio.docx"
doc.save(output)
print(f"Saved: {output}")
